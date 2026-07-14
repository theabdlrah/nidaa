#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate pilot/Nidaa-Overview-FAQ.docx — the external-sharing package.

Three layers (per advisor brief):
  Page 1      : Executive Summary (read first; full doc for reference)
  Current Status box (honest implemented / not-yet)
  Deployed-vs-promised controls note (safety/liability honesty)
  Pages 2-4   : Project Overview (7 subsections)
  FAQ Section  : the skeptical questions a pilot partner asks
Mirrors the style of the repo's make_nidaa_docx.py (teal headings, RTL, Calibri).
Source of truth for depth: pilot/NIDAA-FAQ-OVERVIEW.md (80-Q full FAQ).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\theab\nidaa\pilot\Nidaa-Overview-FAQ.docx"

TEAL  = RGBColor(0x0F, 0x76, 0x6E)
DARK  = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x6B, 0x72, 0x80)
RED   = RGBColor(0xB9, 0x1C, 0x1C)
GREEN = RGBColor(0x15, 0x80, 0x3D)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

sect = doc.sections[0]
sectPr = sect._sectPr
sectPr.append(OxmlElement("w:bidi"))

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = TEAL if level <= 1 else DARK
        r.font.name = "Calibri"
    return h

def para(text, italic=False, bold=False, color=None, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def status_line(label, items, color):
    p = doc.add_paragraph()
    r = p.add_run(label); r.bold = True; r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    for it in items:
        b = doc.add_paragraph(style="List Bullet 2")
        b.add_run(it)
        b.paragraph_format.space_after = Pt(2)

# ============================ PAGE 1: EXEC SUMMARY ============================
title = doc.add_paragraph()
tr = title.add_run("Nidaa (نداء) — Project Overview & FAQ")
tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = TEAL
sub = doc.add_paragraph()
sr = sub.add_run("External-sharing package for NGO coordinators, technical volunteers, "
                "OSM contributors, diaspora organizers, and mutual-aid groups.")
sr.italic = True; sr.font.color.rgb = MUTED; sr.font.size = Pt(10)

para("Read the summary above first; the rest of the document is for reference. "
     "Honesty note: this is an early-stage prototype with no pilot yet — several "
     "safety and privacy controls described below are planned, not built.",
     bold=True, color=DARK, size=10, space_after=8)

bullet(" an offline-first, Arabic-first community board where a local community posts and "
       "browses needs and offers (medical, food, water, shelter), and trusted coordinators "
       "mark entries verified.", bold_lead="What is Nidaa? ")
bullet(" coordination of local help breaks down exactly when connectivity fails. Information "
       "scatters across private chats, gets lost, and cannot be trusted.", bold_lead="What problem does it solve? ")
bullet(" local coordination communities in connectivity-stressed environments (Gaza, West Bank, "
       "Syria named as first targets), grassroots aid groups, diaspora mutual-aid orgs, and the "
       "technical partners who validate the data.", bold_lead="Who is it for? ")
bullet(" the board stays usable with no internet after first visit — posts are saved on the device "
       "and synced when a connection returns. Because the target environments lose connectivity routinely.",
       bold_lead="Why offline-first? ")
bullet(" it is a shared, browsable, verifiable board — not a chat — and verification is role-gated, "
       "audited, and reversible. It complements WhatsApp/Telegram; it does not replace them.",
       bold_lead="What makes it different? ")
bullet(" a board of who-needs-what is a sensitive record. Precise coordinates are optional and flagged "
       "as potentially unsafe; any deployment must weigh surveillance/targeting risk against coordination value.",
       bold_lead="Safety note: ")
bullet(" we are seeking one organization willing to evaluate Nidaa over a 4–8 week period with "
       "approximately 30–100 participants and at least two independent verifiers who jointly own the audit. "
       "Low burden, under one hour of training. This is problem-validation, not a commitment.",
       bold_lead="What we ask of a pilot partner: ")

# ---- Current Status box ----
heading("Current Status", level=1)
para("Nidaa is currently in the pilot-preparation stage. The following is explicit and honest.",
     italic=True, color=MUTED, size=10)

status_line("Implemented (in code, not yet field-tested):", [
    "Offline-first synchronization (posts saved on device, synced when online)",
    "Verification workflow (role-gated; anonymous and ordinary users cannot verify)",
    "Audit logging (every verification action recorded, reversible)",
    "Governance roles defined in code (verifier / admin); operating model PLANNED / unvalidated",
], GREEN)

status_line("Not yet implemented:", [
    "Full identity system (current verification uses server-side tokens, not user accounts)",
    "Large-scale field testing (no pilot has launched)",
    "Production deployment experience (single JSON store; not yet load-tested)",
    "Mesh / peer sync, CRDT conflict resolution, end-to-end encryption, AI assistance",
], RED)

para("Deployed controls vs promised controls (read this before trusting Nidaa with sensitive data):",
     bold=True, color=RED, size=10)
status_line("Actually present today:", [
    "Optional contact fields; no mandatory personal identifiers to post or browse",
    "Single-host JSON store; posts readable by the host operator",
], MUTED)
status_line("Promised, NOT built yet:", [
    "Self-hosting packaging (today you would run on the Nidaa team's host or self-host manually)",
    "End-to-end encryption (data is not E2E-encrypted yet)",
    "Confirmed open-source license (code is public-in-intent; license to be finalized)",
    "External security / ethics review (none yet)",
], RED)
para("Do not treat Nidaa as private or self-controlled until the promised controls ship. A pilot "
     "should not proceed without a do-no-harm review, a basic data-protection assessment, and a "
     "data-processing agreement stating ownership and deletion. (A do-no-harm summary already "
     "exists; the DPIA and agreement are still to be completed.)",
     italic=True, color=MUTED, size=10)

# ---- Architecture flow diagram ----
heading("How information flows (architecture)", level=2)
para("One line of flow. The device is the source of truth; verification and audit sit at the host.",
     italic=True, color=MUTED, size=10)
flow = ["Device\n(browser/PWA)", "Local Storage\n(IndexedDB,\noffline)", "Sync\n(when online)",
        "Host Server\n(JSON store)", "Verification\n(verifier token,\nrole-gated)", "Audit Log\n(recorded,\nreversible)"]
ft = doc.add_table(rows=1, cols=len(flow))
ft.style = "Table Grid"
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, cell in enumerate(ft.rows[0].cells):
    cell.text = ""
    pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for j, line in enumerate(flow[i].split("\n")):
        rr = pr.add_run(line) if j == 0 else pr.add_run("\n" + line)
        rr.font.size = Pt(9); rr.bold = (i in (0, 4, 5))
        if i in (4, 5): rr.font.color.rgb = TEAL
    if i >= 3:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "EAF6F4"); tcPr.append(shd)
arrow = doc.add_paragraph(); arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = arrow.add_run("      →      →      →      →      →      ")
ar.font.size = Pt(12); ar.font.color.rgb = MUTED
para("Offline posting and browsing work at the Device + Local Storage stage with no connection. "
     "Sync, Verification, and Audit require a connection to the Host — so during a network shutdown "
     "the 'verified' state cannot be updated. Nidaa preserves unverified local posts; it does NOT "
     "solve trust under a total blackout. Mesh sync (device-to-device, no host) is planned, not built.",
     size=10, color=DARK)

doc.add_page_break()

# ============================ PAGES 2-4: OVERVIEW ============================
heading("Project Overview", level=1)

heading("Why Nidaa? — comparison", level=2)
para("Nidaa is not a messaging app. The differences that matter for coordination:", size=10, color=MUTED)
cmp_rows = [
    ("Capability", "WhatsApp", "Telegram", "Nidaa"),
    ("Post offline (no connection)", "No", "No", "Yes (device-first)"),
    ("Browse history offline", "Yes (local)", "Yes (local)", "Yes"),
    ("Shared browsable board", "No", "Limited (channels)", "Yes"),
    ("Role-gated verification", "No", "No", "Yes (verifier/admin)"),
    ("Audit trail of verification", "No", "No", "Yes (recorded, reversible)"),
    ("End-to-end encryption", "Yes (default)", "Optional", "Not yet (planned)"),
    ("Self-hostable", "No", "No", "Intended (packaging planned)"),
]
ct = doc.add_table(rows=len(cmp_rows), cols=4)
ct.style = "Table Grid"
for ri, row in enumerate(cmp_rows):
    for ci, val in enumerate(row):
        c = ct.rows[ri].cells[ci]
        c.text = ""
        pp = c.paragraphs[0]
        rn = pp.add_run(val); rn.font.size = Pt(9.5)
        if ri == 0:
            rn.bold = True; rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tcPr = c._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "0F766E"); tcPr.append(shd)
        elif ci == 3:
            rn.bold = True; rn.font.color.rgb = TEAL
para("Honest note: WhatsApp/Telegram are far more mature and field-tested. Nidaa's only claimed "
     "differentiators are offline-first posting (device is source of truth) and audited, reversible "
     "verification. Whether those matter in practice is unvalidated. Nidaa today is a resilient "
     "LOCAL list per community; cross-community coordination depends on mesh/federation, which is "
     "planned, not built.", italic=True, color=MUTED, size=9.5)

heading("1. The coordination problem", level=2)
para("In crises, the channels people rely on — messaging groups, central servers, social platforms — "
     "become unreliable at the moment needs spike. Information scatters across closed chats newcomers "
     "cannot see, is lost when phones die or accounts are banned, and cannot be trusted because anyone "
     "can post anything. Nidaa targets the information-coordination layer only: it does not solve "
     "logistics, physical delivery, funding, or personal security. Important: if a community's real "
     "bottleneck is physical supply (resources, not information), Nidaa does not address it.")

heading("2. How Nidaa works", level=2)
bullet("A Next.js web app (installable PWA) that runs in a browser on a phone.")
bullet("Posts (needs/offers) are written to the device first (IndexedDB), so posting works with no connection.")
bullet("When a connection is available, queued posts sync to a server, which stores them in a JSON file.")
bullet("The board displays entries by type, category, region, and verification status.")
bullet("A trusted verifier can mark an entry verified through a protected endpoint; the action is audited and reversible.")
bullet("The map seeds facility locations from humanitarian datasets (HDX / OpenStreetMap) — see the staleness caveat below.")

heading("3. Verification model", level=2)
para("Verification is institutional, not algorithmic. A reader trusts a “verified” mark because a party "
     "the deploying organization designated as a verifier made it. The endpoint rejects anonymous and "
     "ordinary-user tokens (returns 401); only verifier/admin tokens succeed (tested). Every action is "
     "written to an audit log (entry, actor role, prior/new state, timestamp) and is reversible. Nidaa "
     "cannot make information true — it can only make the act of verification accountable.")
para("Limitation: verification requires a connection to the Host. During a network shutdown the "
     "'verified' state cannot be updated, so trust signals are stale precisely when crisis peaks. "
     "Mitigation requires at least two independent verifiers per deployment and an out-of-band override "
     "the org defines; a single verifier has no automated fallback.", color=DARK, size=10)

heading("4. Offline-first architecture", level=2)
para("The device is the source of truth. New posts are saved locally with a “pending” flag, then synced "
     "when online. Browsing works fully offline. Current limitations: sync is client→server only (no "
     "mesh), conflict resolution is last-write-wins (no CRDT), and there is no end-to-end encryption yet. "
     "Silent data loss is possible if two offline devices edit the same entry; this is a known safety "
     "concern for life-saving posts and is not yet mitigated in the UI.")

heading("5. Governance model", level=2)
para("The deploying organization defines verifier roles, owns the audit, and handles disputes. The "
     "project maintainer governs the software; pilot governance sits with the partner. Verification is "
     "transparent (auditable) and reversible so abuse can be detected and corrected — but trust in a "
     "verifier is organizational, not technical, and the trust-bootstrap (who legitimately designates "
     "the verifier) is an open, unresolved problem.")
para("Responsibilities in a pilot:", bold=True, size=10)
rt = doc.add_table(rows=3, cols=2); rt.style = "Table Grid"
resp = [("Partner provides", "Nidaa team provides"),
        ("≥2 independent verifiers + audit ownership", "Hosting + sync infrastructure"),
        ("Local governance, dispute handling, staff time", "Software, metrics, co-authored field report")]
for ri, row in enumerate(resp):
    for ci, val in enumerate(row):
        c = rt.rows[ri].cells[ci]; c.text = ""
        pp = c.paragraphs[0]; rn = pp.add_run(val); rn.font.size = Pt(9.5)
        if ri == 0:
            rn.bold = True; rn.font.color.rgb = TEAL

heading("6. Known limitations", level=2)
bullet("Single JSON store — not built for high concurrency or large scale.")
bullet("No real user identity yet; token leakage would compromise verification.")
bullet("Naive conflict resolution (last-write-wins, silent loss); no mesh / offline peer sync.")
bullet("No end-to-end encryption; privacy architecture not yet externally audited.")
bullet("Map facility data is seeded from HDX/OSM without a snapshot date or update path; in active "
       "conflict it can be outdated — treat map facilities as indicative until validated.")
bullet("In some sandboxed environments the board may render empty if it cannot reach the API; a live "
       "deployment avoids this.")
bullet("Depends on at least one host (mesh is planned, not built).")
bullet("Scope: Nidaa is a resilient local list per community; cross-community coordination needs "
       "mesh/federation (planned, not built).")

heading("7. Pilot structure", level=2)
para("Intended: a low-risk, instrumented pilot with one reachable coordination community "
     "(diaspora mutual-aid orgs are the current outreach target; Gaza NGOs are deferred as currently "
     "unreachable). 4–8 weeks, low burden, under one hour of training. The partner defines verifier "
     "roles and owns the audit; the Nidaa team provides the host and metrics and co-authors a field "
     "report. Success = evidence the board helped coordination; failure = documented lessons. No pilot "
     "has launched; outreach is prepared but not yet sent. Before any pilot, the safety/liability items "
     "above (do-no-harm, DPIA, data agreement) must be in place.")

doc.add_page_break()

# ============================ FAQ SECTION ============================
heading("Frequently Asked Questions", level=1)
para("The most decision-relevant questions a skeptical partner asks. The full 80-question FAQ is in "
     "NIDAA-FAQ-OVERVIEW.md.", italic=True, color=MUTED, size=10)

FAQ = [
    ("Is this a social network?",
     "No. It is a shared, browsable board for needs and offers — not a feed, profile, or chat network. "
     "But we do not dodge the surveillance question: a board of who-needs-what is a sensitive record, "
     "and we treat that as a primary risk, not a side note."),
    ("Is this encrypted?",
     "Not end-to-end yet. Posts are stored on the device and on the host server in a readable form. "
     "E2E is planned but not built. Self-hosting is intended but not yet packaged. Do not treat Nidaa "
     "as private until those ship."),
    ("What happens when connectivity is lost?",
     "Posting and browsing keep working on the device. Queued posts sync automatically when a connection "
     "returns. Cross-device sync without any host requires mesh networking, which is planned, not built."),
    ("Who can verify information?",
     "Only designated verifiers or admins with a token. Anonymous users and ordinary users cannot — the "
     "endpoint returns 401 for them. This was tested."),
    ("What if a verifier is compromised?",
     "Every verification action is audited (actor, prior/new state, timestamp) and reversible, so another "
     "verifier can undo it — provided at least two independent verifiers exist. With a single verifier "
     "there is no automated fallback; the org must define an out-of-band override. Trust in the verifier "
     "is the organization's responsibility, not a software guarantee."),
    ("How is misinformation handled?",
     "Primarily by human verification and by visually distinguishing verified from unverified entries. There "
     "is no automated false-information detection and no moderation AI yet. Reporting and community "
     "escalation are planned, not built."),
    ("Why not just use WhatsApp?",
     "WhatsApp is private messaging; posts there are not browsable, persist per-account, and anyone can claim "
     "anything. Nidaa is a shared board that survives disconnection and is verifiable. It can complement "
     "WhatsApp (e.g. receive posts from a group), not replace it. We acknowledge WhatsApp is more mature "
     "and has network effects we do not."),
    ("Why not just use Telegram?",
     "Same reasoning as WhatsApp: Telegram is a messaging platform without a verifiable, offline-first, "
     "browsable coordination record. Nidaa's differentiators are the device-is-source-of-truth design and "
     "audited, reversible verification."),
    ("Who owns the data?",
     "Intended: the deploying organization, via self-hosting and export. Export UI and self-host packaging "
     "are not yet built; today data would sit on the Nidaa team's host. A pilot data-processing agreement "
     "stating ownership and deletion is still to be completed."),
    ("What stage is the project in?",
     "Active development, prototype stage, pilot-preparation. Core coordination, offline-first storage, and "
     "role-gated verification with audit exist. Identity, mesh, CRDT, E2E, AI assistance, and any real pilot do not."),
    ("What has and has not been built yet?",
     "Built: offline-first PWA, local device store, server entries API, role-gated verification + audit "
     "(reversible), Arabic-first RTL UI, HDX/OSM map seed, ~10k illustrative entries. Not built: real identity, "
     "mesh/CRDT sync, end-to-end encryption, automated moderation, self-host packaging, external security/"
     "ethics review, and any pilot."),
    ("What support is required from a pilot partner?",
     "At least two independent verifiers, 30–100 participants, low burden, 4–8 weeks, under one hour of "
     "training, and ownership of local governance/audit. The Nidaa team provides hosting, software, and metrics."),
    ("Is the seed data real operational data?",
     "No. The sample entries are illustrative (Syria cities) and explicitly not real operational data. The "
     "map facilities are seeded from HDX/OSM but carry no snapshot date — treat them as indicative, not current."),
    ("Can it run with no internet at all?",
     "Local posting and browsing: yes. Cross-device synchronization: needs a host server today, or mesh "
     "(planned, not built) in the future."),
    ("What about the targeting / surveillance risk?",
     "It is real. A geolocated board of needs can be weaponized. Mitigations today are limited: optional "
     "precise coordinates (flagged as potentially unsafe), no mandatory personal data, and self-hosting "
     "(planned). We recommend deployments minimize precise coordinates and weigh this risk explicitly "
     "before launch. This is the primary ethical risk, not a footnote."),
    ("What is the single biggest risk to Nidaa?",
     "Two kinds, ranked by harm to affected people: (1) physical/safety harm if the board is misused for "
     "targeting or misinformation; (2) no real-world validation — building infrastructure communities do not "
     "adopt. The pilot is the instrument for answering the second; the safety items above must be addressed "
     "before the first."),
]

for q, a in FAQ:
    p = doc.add_paragraph()
    r = p.add_run("Q: " + q); r.bold = True; r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(1)
    ap = doc.add_paragraph()
    ar = ap.add_run(a); ar.font.size = Pt(10.5)
    ap.paragraph_format.space_after = Pt(8)

para("— End of overview. For the complete 80-question FAQ, implementation status, and ethical-risk "
     "detail, see NIDAA-FAQ-OVERVIEW.md in the same package.", italic=True, color=MUTED, size=9)

doc.save(OUT)
print("WROTE:", OUT)
