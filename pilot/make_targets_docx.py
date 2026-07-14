#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate C:\\Users\\theab\\Desktop\\Nidaa-Outreach-Targets.docx
Targets + fill-in drafts for the user to fine-tune and send.
Style mirrors pilot/make_faq_docx.py (teal headings, Calibri).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"C:\Users\theab\Desktop\Nidaa-Outreach-Targets.docx"
TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x6B, 0x72, 0x80)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)

def H(text, lvl=1):
    h = doc.add_heading(text, level=lvl)
    for r in h.runs:
        r.font.color.rgb = TEAL if lvl <= 1 else DARK
        r.font.name = "Calibri"
    return h

def P(text, italic=False, bold=False, color=None, size=11, space=6):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space)
    return p

def B(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text); p.paragraph_format.space_after = Pt(3)
    return p

# Title
t = doc.add_paragraph(); tr = t.add_run("Nidaa — Outreach Targets & Drafts")
tr.bold = True; tr.font.size = Pt(18); tr.font.color.rgb = TEAL
P("Problem-discovery framing only — no pilot, no demo, no commitment. "
  "Fine-tune [name] and [PACKAGE LINK], then send. Fill [PACKAGE LINK] with one "
  "shared Drive/Google Folder of the pilot/ directory (or the docx).",
  italic=True, color=MUTED, size=10)

H("Tier 1 — Diaspora mutual-aid orgs (primary, reachable)", 1)

def target(name, who, why, where, subject, body):
    H(name, 2)
    B("Who: " + who)
    B("Why: " + why)
    B("Where to reach: " + where)
    P("Draft:", bold=True, size=10, space=2)
    p = doc.add_paragraph(); r = p.add_run("Subject: " + subject)
    r.bold = True; r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    pb = doc.add_paragraph(); rb = pb.add_run(body)
    rb.font.size = Pt(10.5); pb.paragraph_format.space_after = Pt(8)

target("1. The Sameer Project",
       "diaspora-led mutual-aid coordinating Gaza medical/relief + local partners.",
       "Exactly the connectivity-stressed coordination environment Nidaa targets.",
       "Instagram @thesameerproject / site contact form / public email. Find a programs or coordination lead.",
       "Researching coordination gaps in Gaza aid work — could I get your perspective?",
       "Hi [name],\n\nI'm Abdul Rahman, a student researching how aid coordination actually works in "
       "connectivity-stressed environments like Gaza. I'm NOT reaching out to pitch a tool — I'm trying "
       "to understand whether the coordination problems I think exist are real and urgent, or whether "
       "I've got the wrong end of the stick.\n\nIf you have 15–30 minutes in the next couple of weeks, "
       "I'd value your honest perspective on how your team coordinates needs and offers today, what "
       "breaks when connectivity drops, and where the real friction is. I'd especially like to be "
       "corrected if my assumptions are off — I suspect a lot of what I've assumed is wrong.\n\nNo ask "
       "beyond a conversation. Happy to share what I've found so far afterward if useful. Worth a "
       "short call?\n\n(Background if you want it: [PACKAGE LINK])")

target("2. HEAL Palestine",
       "diaspora-led org focused on children/families in Gaza.",
       "Coordinates fragmented, low-connectivity work; good signal on matching vs supply.",
       "Site contact / Instagram @heal.palestine / public email.",
       "Researching coordination friction in diaspora aid orgs — 20 min of your perspective?",
       "Hi [name],\n\nI'm Abdul Rahman, researching how diaspora-led aid orgs coordinate fragmented, "
       "low-connectivity work like yours. I'm at the stage where I need to be wrong in front of people "
       "who know, rather than wrong in private.\n\nI'd value 15–30 minutes on: how coordination actually "
       "happens in your work, what fails during connectivity collapses, and whether the \"information "
       "coordination\" problem I'm fixated on is even the binding constraint — or whether supply, trust, "
       "or something else dominates. If Nidaa (what I've built so far) is aiming at the wrong problem, "
       "I'd rather hear that now.\n\nJust a conversation, nothing expected in return. Could we find 20 "
       "minutes?\n\n(Background if useful: [PACKAGE LINK])")

target("3. Gaza Soup Kitchen (or similar grassroots group)",
       "grassroots local aid coordination inside/around Gaza.",
       "Ground-level view of what actually breaks; strongest reality-check.",
       "Instagram presence / mutual-aid network refs. Confirm the active handle before sending — these change.",
       "Trying to understand aid coordination under connectivity stress — your take?",
       "Hi [name],\n\nAbdul Rahman here. I'm doing informal research on how grassroots aid groups coordinate "
       "local needs and resources when the internet is unreliable or cut — and I keep finding that my "
       "guesses about what's hard don't match reality on the ground. I'd rather learn from people doing "
       "the work than keep guessing.\n\nCould I get 15–30 minutes of your perspective? I'm specifically "
       "trying to find out what actually breaks in current workflows, whether existing tools (WhatsApp, "
       "Telegram, spreadsheets) are good enough, and whether a coordination problem is even a priority "
       "versus, say, a resources problem. I'm wide open to being told Nidaa — the thing I've been "
       "building — is solving the wrong problem.\n\nNo commitment, no pitch. Just a conversation. Open to it?")

H("Tier 2 — Technical / data validator (parallel track)", 1)
target("4. HOTOSM (Humanitarian OpenStreetMap Team)",
       "Maintains the HDX/OSM facility data Nidaa's map seeds from.",
       "Best external eye on data-integrity assumptions; may point to ground orgs.",
       "HOT Slack / community forum / staff emails on hotosm.org.",
       "How do you handle facility-data freshness in active conflict? — research question",
       "Hi [name] at HOT,\n\nAbdul Rahman here. I'm researching coordination challenges in crisis "
       "environments and keep running into a data problem I don't know the right answer to: when you "
       "seed a coordination board with HDX/HOT OSM facility data in an active conflict, how do you "
       "handle staleness and provenance so you don't send people to a clinic that no longer exists?\n\n"
       "I'm NOT proposing a deployment or a pilot. I'd genuinely value 15–30 minutes of your perspective "
       "on where my assumptions about data integrity and offline-first coordination fall apart — and "
       "which of my \"obvious\" ideas are actually naive. Happy to share what I've built so far if it "
       "helps you react, but the conversation is the point, not the tool.\n\nOpen to a chat?")

H("Sending order (suggested)", 1)
B("1. Sameer Project — highest-fit, most likely to reply to research framing.")
B("2. HEAL Palestine — parallel diaspora signal.")
B("3. HOTOSM — technical track; may yield a ground-org referral.")
B("4. Gaza Soup Kitchen — grassroots; hardest signal, send once you've confirmed handle.")
P("Send 3–5. After ANY reply, complete assumption-log/<ORG>-<DATE>.md and update the tracker "
  "BEFORE next steps. No \"they liked it\" summaries.", italic=True, color=MUTED, size=10)

doc.save(OUT)
print("WROTE:", OUT)
