from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# RTL on the whole document (Arabic + Latin mix)
sect = doc.sections[0]
sectPr = sect._sectPr
bidi = OxmlElement("w:bidi")
sectPr.append(bidi)

TEAL = RGBColor(0x0F, 0x76, 0x6E)
DARK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = TEAL if level <= 1 else DARK
        run.font.name = "Calibri"
    return h


def add_para(text, italic=False, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    # RTL run property so Arabic renders correctly
    rPr = r._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    return p


def add_bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        rPr = r._element.get_or_add_rPr()
        rPr.append(OxmlElement("w:rtl"))


# ===== Title =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Nidaa (نداء)")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = TEAL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("An Offline-First Community Coordination Board for Conflict- and Crisis-Affected Populations")
sr.font.size = Pt(13)
sr.italic = True
sr.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("A project journal — abstract through purpose, uniqueness, and field intent\n"
                  "Author: Abdul Rahman (@theabdlrah)  ·  Repository: github.com/theabdlrah/nidaa\n"
                  "Status: Working prototype (offline-first PWA, Arabic-RTL)  ·  Also applicable to: Gaza")
mr.font.size = Pt(9.5)
mr.font.color.rgb = MUTED
mrPr = mr._element.get_or_add_rPr()
mrPr.append(OxmlElement("w:rtl"))

doc.add_paragraph()

# ===== 1. Abstract =====
add_heading("1. Abstract", 1)
add_para(
    "Nidaa (نداء — \"call\" or \"appeal\") is an offline-first, Arabic-first community coordination "
    "board designed for populations living under constrained or destroyed connectivity. Unlike "
    "conventional web platforms that assume a stable internet connection, Nidaa treats the device "
    "as the source of truth: a new post is written to local storage first and synchronized to a "
    "server only when a connection becomes available. The application shell is cached by a service "
    "worker, so the board remains fully usable — browse, read, and compose — with zero connectivity "
    "after the first visit. Gaza is the primary deployment target; the same codebase serves other "
    "low-connectivity crisis regions (e.g. Syria) as region-scoped instances."
)
add_para(
    "The problem Nidaa addresses is not \"lack of an app\"; it is that the people who most need "
    "coordination (displaced families, local committees, volunteers, under-resourced clinics) are "
    "precisely the ones for whom online-only tools fail. By inverting the normal model — offline by "
    "default, online when possible — Nidaa stays operational through internet shutdowns, damaged "
    "infrastructure, and slow or expensive mobile data."
)

# ===== 2. Context & Motivation =====
add_heading("2. Context & Motivation", 1)
add_para(
    "The design constraints are defined directly by current conditions in Gaza and Syria (2025). "
    "Gaza is the focus of deployment: connectivity collapses under conflict, aid-distribution sites "
    "have been struck, and the population is almost entirely dependent on intermittent, contested "
    "connectivity and humanitarian assistance. Syria provides the documented baseline below; the "
    "constraints are shared in kind:"
)
add_bullets([
    "~64% of the population is effectively offline; internet penetration is only ~35.8% (9.01M users of 25.2M people).",
    "Median mobile download speed is ~12.7 Mbps, intermittent and costly.",
    "Roughly half of all hospitals, schools, and water systems are severely damaged or destroyed (WHO, 2025).",
    "A new interim government and a national digital-transformation strategy to 2030 signal appetite for digital public services — but the field reality remains low-connectivity and device-constrained.",
])
add_para(
    "These are not edge cases; they are the baseline. Any tool that cannot function under them is, in "
    "practice, unavailable to the people it claims to serve. The same baseline describes Gaza, where "
    "connectivity is repeatedly severed, civilian infrastructure is devastated, and coordination of aid, "
    "medical care, and family reconnection happens under extreme disruption. Nidaa's architecture is "
    "purpose-built for exactly this class of environment, which is why the design transfers directly."
)

# ===== 3. Purpose =====
add_heading("3. Purpose", 1)
add_para(
    "Nidaa exists to let a community self-coordinate local needs and resources without depending on "
    "central infrastructure or constant connectivity. Concretely, it enables:"
)
add_bullets([
    "Discovery — browsing community needs and offers (medical, food, water, shelter, education, transport, other) by city.",
    "Local posting under any condition — a volunteer or committee member can post while offline; the entry is queued on-device and pushed automatically on reconnect.",
    "Trust signals — entries can be marked verified by trusted NGOs/admins, so unverified community reports are visibly distinguished from confirmed ones.",
    "Resilience by default — online/offline and pending-sync status are always shown; no action is ever lost because the network dropped.",
])
add_para(
    "The intent is not to replace national humanitarian information systems (e.g. Sahana EDEN, UN "
    "coordination platforms). Those serve a different scale. Nidaa is the neighborhood-level board — "
    "lightweight, owned by the community, and survivable when everything else goes dark."
)

# ===== 4. Why Unique =====
add_heading("4. Why This Application Is Unique", 1)
add_para(
    "Most existing tools in this space fall into two camps, and Nidaa sits deliberately between and "
    "apart from them:"
)
add_bullets([
    "Heavy national IMS platforms (Sahana EDEN, HDX pipelines) are powerful but require training, connectivity, and institutional backing a local committee does not have.",
    "Offline first-aid / EMR tools (ReliefGPT, QuickAid, Project Buendia) solve specific medical scenarios and carry liability constraints.",
])
add_para("Nidaa's distinct contribution is the offline-first community board — a general, low-friction coordination surface that:")
add_bullets([
    "Runs on cheap Android devices as an installable PWA — no app-store friction, no native build.",
    "Is Arabic-RTL first, not a translated afterthought.",
    "Guarantees no data loss on disconnect via local-first write + retry sync.",
    "Uses zero native dependencies on the server (a JSON store) so it deploys anywhere Node runs — and the API shape already supports upgrading to Postgres.",
    "Is buildable and auditable by a small team, which matters for trust in sensitive contexts.",
])
add_para(
    "The uniqueness is architectural, not feature-count: it assumes the network will fail, and is "
    "correct anyway."
)

# ===== 5. What We Are Looking For =====
add_heading("5. What We Are Looking For", 1)
add_para(
    "Nidaa is a prototype, not a deployed service. To move from prototype to something a community can "
    "rely on, we are explicitly looking for:"
)
add_bullets([
    "Verified data sources. Seed data today is illustrative only. Real value requires NGO- and committee-sourced, fact-checked entries — not scraped or assumed information.",
    "Authentication & role-gated verification. The /api/verify route is open in the prototype. In production it must be behind auth so only trusted actors can mark entries verified (prevents misinformation / manipulation).",
    "Offline maps. Coordinates are stored; the next layer is Leaflet with offline-capable tiles so locations render without a live connection.",
    "Field testing. The only real validation is use by an actual committee or volunteer network under real conditions — including Gaza.",
    "Privacy & safeguarding. Especially if extended toward sensitive categories (e.g. family reconnection), strong privacy and consent controls are mandatory, not optional.",
])
add_para(
    "We are not looking to build a surveillance-adjacent or data-extractive tool. The model is "
    "local-first, community-owned, and resilient against interception."
)

# ===== 6. Primary Deployment: Gaza =====
add_heading("6. Primary Deployment: Gaza", 1)
add_para(
    "Gaza is the primary deployment target for Nidaa, not a sibling case. The constraints the design "
    "was built around — severed or contested connectivity, devastated infrastructure, acute need for "
    "local coordination of aid, medical care, water, and shelter, and a predominantly Arabic-speaking "
    "population — are present there in their most acute form. Crucially, the threat model (Section 12) "
    "is shaped by Gaza's realities: aid sites have been struck, so location precision and PII exposure "
    "are live risks, not hypotheticals — which is exactly why geo-indistinguishability (10.6) and "
    "no-mandatory-PII (12.2) are core, not optional."
)
add_para("Deploying to Gaza requires no architectural change; it requires:")
add_bullets([
    "Gaza-specific seed/verified data from trusted local actors (already importable: State of Palestine "
    "health + education facilities from HDX/HOT OSM — clipped to the Gaza Strip (1,315 facilities) and "
    "West Bank (3,397) via the Gaza Strip Municipal Boundaries; Syria is a separate secondary region).",
    "A region filter in the UI (All / Gaza Strip / West Bank / Syria) so the same deployment scopes to "
    "either Palestinian territory or Syria without code changes.",
    "A Gaza-facing trusted-verifier set: UNRWA, Palestinian Red Crescent, and vetted local Gaza NGOs pre-seeded as verifiers (Phase 3), with offline-checkable signed credentials (10.6).",
])
add_para(
    "Current-state note (2026-07-14): the import pipeline (scripts/import-hdx.mjs) now strips any "
    "illustrative 'demo' rows on import (honesty fix B, commit 09a8573), so the live board holds only "
    "real HDX/HOT OSM facilities — 10,382 entries, 0 fabricated. The clipping counts above (Gaza 1,315 / "
    "West Bank 3,397) describe the import design and remain accurate as the real-facility split.",
    italic=True, color=MUTED, size=10,
)
add_para(
    "Syria remains a secondary, region-scoped instance of the same codebase — the offline-first core "
    "is shared; only data and localization differ. This generality is deliberate: the authors are not "
    "from either region, so Nidaa is built as a generalizable coordination primitive, Gaza-first."
)

# ===== 7. Limitations =====
add_heading("7. Honest Limitations (stated up front)", 1)
add_bullets([
    "Board data is REAL HDX/HOT OSM facility data: 10,382 verified facilities imported (Gaza/West Bank "
    "primary, Syria secondary). It is curated OSM data, not live operational feeds — still requires "
    "local-actor validation before field use. NOTE (honesty fix B, 2026-07-14, commit 09a8573): the "
    "board contains 0 fabricated entries; the earlier illustrative 'demo seed' rows were stripped so a "
    "trust-and-verification tool never ships made-up data. See Build Log 14.5 correction and 14.12.",
    "The verification endpoint is open in this build and must be secured (Phase 3).",
    "The server store is a JSON file (chosen for zero-dependency portability); scale demands SQLite/Postgres (Phase 4).",
    "Offline maps are implemented (Phase 2) with IndexedDB tile cache; tiles cover visited areas only to stay light on cheap devices.",
    "This is a portfolio / learning prototype demonstrating offline-first architecture for constrained-connectivity humanitarian contexts, Gaza-first.",
])

# ===== 8. Phased Build Plan =====
add_heading("8. Phased Build Plan", 1)
add_para(
    "The roadmap from prototype to a field-usable tool. Each phase ends with a verification gate "
    "that must pass before the next begins. Research basis: HDX (236 Syria datasets, including 3,060 "
    "real health facilities as GeoJSON), ReliefWeb API v2 (live, free), Auth.js v5, better-sqlite3, "
    "and Leaflet offline. Guiding principle: the network will fail, and the tool must be correct anyway."
)

add_heading("Phase 0 — Harden & Verify Baseline (prerequisite)", 2)
add_para(
    "Lock the prototype with a real API test (GET returns array, POST creates, verify flips flag). "
    "Gate: build green; API smoke test passes; Arabic round-trip OK.", italic=False)

add_heading("Phase 1 — Real Seed Data from HDX  (highest impact, lowest effort)", 2)
add_para(
    "Replace illustrative seeds with verified, real facility data. A scripts/import-hdx.mjs downloader "
    "maps OSM amenities to Nidaa categories and upserts them as verified 'offer' entries. Gaza-first: "
    "State of Palestine (Gaza/West Bank) health (2,053) + education (2,659) facilities are PRIMARY; "
    "Syria health (2,925) + education (2,745) are a SECONDARY region instance. A 'region' field "
    "(gza | syr) enables region-scoped deployments sharing the offline-first core."
)
add_para("Gate: importer run populates the board; a known Gaza facility is present and searchable by city (e.g. a مستشفى/عيادة entry in Gaza City).", italic=True, color=MUTED, size=10)

add_heading("Phase 2 — Offline Maps (Leaflet + cached tiles)", 2)
add_para(
    "Add Leaflet (client-only, dynamic import) and cache viewed tiles in IndexedDB when online, "
    "serving them offline — only for visited areas to stay light on cheap devices. List/Map toggle; "
    "tap a card to center the map. Gaza reuses the same component."
)
add_para("Gate: map renders online; after visiting an area, going offline still shows tiles with no console errors.", italic=True, color=MUTED, size=10)

add_heading("Phase 3 — Auth + Role-Gated Verification (closes security gap)", 2)
add_para(
    "Adopt Auth.js v5 with roles (individual | volunteer | ngo | admin). POST /api/verify now "
    "requires an ngo|admin session (401 otherwise) — closing the open verify route flagged in Section 7. "
    "Local posts from unverified users stay 'unverified' until an NGO confirms. For Gaza, the trusted-"
    "verifier set is PRE-SEEDED as a governance choice (not invented by code): UNRWA, the Palestinian "
    "Red Crescent, and vetted local Gaza NGOs, each holding an offline-checkable signed credential "
    "(DID/VC, Section 10.6). New NGOs are verified by existing trusted NGOs (web-of-trust), never by a "
    "lone central server. An NGO review queue lets trusted actors verify or reject with a reason, and "
    "all verify actions are audit-logged + reversible (Section 12.2)."
)
add_para("Gate: unauthenticated verify → 401; ngo session → 200 and flag flips; individual → 401.", italic=True, color=MUTED, size=10)

add_heading("Phase 4 — Server Store Upgrade: JSON → SQLite", 2)
add_para(
    "Introduce better-sqlite3 (synchronous, zero-config) behind lib/store.ts with identical function "
    "signatures. Handles thousands of HDX rows and concurrent writes safely. Postgres only if "
    "multi-instance/HA is later required. Existing API shape is unchanged."
)
add_para("Gate: all Phase 1–3 behaviors pass against SQLite; bulk import of 3,060 rows completes without corruption; concurrent POSTs lose no writes.", italic=True, color=MUTED, size=10)

add_heading("Phase 5 — Sync Conflict Resolution (hard offline-first problem)", 2)
add_para(
    "Handle multiple volunteers editing the same entry offline. Server assigns the canonical id on "
    "first sync; merge rule uses updatedAt (wins), and if a client edited after its last pull the "
    "entry is flagged 'conflict' with both versions surfaced to an NGO reviewer. A sync log records "
    "each push/pull for audit."
)
add_para("Gate: two simulated offline edits to one entry are both preserved, conflict flagged, and resolvable with no data loss.", italic=True, color=MUTED, size=10)

add_heading("Stretch (post-Phase 5, optional)", 2)
add_bullets([
    "WhatsApp/SMS bridge: post needs via message (Twilio / WhatsApp Business API) — reaches the ~64% who will not install an app; handle PII carefully.",
    "Multi-region deploy: same code, region=syr|gza scoping, shared offline-first core.",
    "ReliefWeb context cards: pull v2 reports for a city as situational context.",
])

add_para(
    "Sequencing: P0 → P1 real data → P2 maps → P3 auth → P4 SQLite → P5 conflicts → stretch. The "
    "trio that moves Nidaa from fake-data prototype to real, safe, scalable tool is P1 + P3 + P4.",
    bold=True,
)

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    # shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)
    return p


# ===== 10. Architecture Notes: mechanism-level improvements =====
add_heading("10. Architecture Notes: Mechanism-Level Improvements", 1)
add_para(
    "Five mechanism-level upgrades move Nidaa from offline-tolerant toward offline-native. They are "
    "listed roughly in order of field impact. Items 1 and 4 include pseudocode sketches.",
)

add_heading("10.1 Device-to-device sync (the biggest gap)", 2)
add_para(
    "Current sync is device↔server only — it assumes the server is reachable when connectivity "
    "returns. But the stated worst case is a total internet shutdown, where the server is unreachable "
    "for everyone at once. Local mesh sync lets devices in the same room, building, or camp exchange "
    "posts directly over local WiFi/hotspot (WebRTC) or Bluetooth, so the community board keeps "
    "updating with zero internet. This is the difference between 'offline-tolerant' and 'offline-native.'"
)
add_para("Sketch — WebRTC data-channel gossip between peers on a local network:", bold=True)
add_code(
    "// each device runs a lightweight gossip peer\n"
    "const peer = new GossipPeer({ id: localClientId });\n"
    "\n"
    "peer.on('entry', (remote) => {\n"
    "  if (alreadyHave(remote.clientId)) return;          // idempotent\n"
    "  if (isNewer(remote, localCopy(remote.clientId))) {\n"
    "    upsertLocal(remote);                             // write to IndexedDB\n"
    "    peer.gossip(remote, { exclude: remote.from });   // forward once\n"
    "  }\n"
    "});\n"
    "\n"
    "// on reconnect, push the merged local set to the server\n"
    "window.addEventListener('online', () => pushPending());"
)

add_heading("10.2 Idempotency on retry", 2)
add_para(
    "When a queued post finally syncs and the request partially fails (times out after the server "
    "received it but before the client got the ack), a naive retry duplicates the entry. Every "
    "local-first post already carries a client-generated UUID at creation; the server must dedupe on "
    "that key (upsert by clientId). This is already implemented in lib/store.ts — it is called out "
    "here because it is the invariant that makes 10.1 and 10.5 safe under retries."
)

add_heading("10.3 Sync prioritization, not just queuing", 2)
add_para(
    "With 12.7 Mbps median and costly data, the sync queue should be priority-ordered, not FIFO: "
    "medical/urgent categories sync first; general posts batch later. Compress payloads and defer any "
    "images until the device is on WiFi (or drop them in a true blackout)."
)
add_para("Sketch — priority-ordered drain:", bold=True)
add_code(
    "const PRIORITY = { medical: 0, water: 1, shelter: 2, food: 3, education: 4, other: 9 };\n"
    "async function drain() {\n"
    "  const q = (await pendingLocal()).sort(\n"
    "    (a, b) => PRIORITY[a.category] - PRIORITY[b.category]);\n"
    "  for (const e of q) {\n"
    "    if (e.hasImage && !onWifi()) continue;            // defer heavy media\n"
    "    await postCompressed(e);                          // gzip before send\n"
    "  }\n"
    "}"
)

add_heading("10.4 Field-level conflict merge, not whole-entry flagging", 2)
add_para(
    "The Phase 5 plan flags the whole entry as 'conflict' if two people edit offline. For simple "
    "cases (one adds a phone number, another updates status) that is overkill. A field-level "
    "three-way merge resolves most conflicts automatically and only escalates true clashes (same "
    "field, different values) to an NGO reviewer."
)
add_para("Sketch — three-way field merge:", bold=True)
add_code(
    "function merge(base, local, remote) {\n"
    "  const out = { ...base };\n"
    "  const clashes = [];\n"
    "  for (const k of keys(local, remote)) {\n"
    "    if (local[k] === remote[k]) { out[k] = local[k]; }\n"
    "    else if (local[k] === base[k]) { out[k] = remote[k]; }   // only remote changed\n"
    "    else if (remote[k] === base[k]) { out[k] = local[k]; }   // only local changed\n"
    "    else clashes.push(k);                                    // both diverged -> human\n"
    "  }\n"
    "  out.clashes = clashes;                                     // [] => auto-resolved\n"
    "  return out;\n"
    "}"
)

add_heading("10.5 Decentralized trust, not single-gate verification", 2)
add_para(
    "Verification is currently binary: an NGO/admin marks an entry verified. In a real blackout the "
    "NGO account may be unreachable too — a single point of failure for trust. A lightweight "
    "corroboration model fixes this: N independent community members confirming an entry raises its "
    "trust tier even before an NGO acts. Trust becomes a gradient, not a gate, so the board stays "
    "useful when central authorities are dark."
)
add_para("Sketch — corroboration tiers:", bold=True)
add_code(
    "function trustTier(entry) {\n"
    "  const n = entry.corroborations.length;        // distinct community confirmations\n"
    "  if (entry.verifiedByNgo) return 'VERIFIED';\n"
    "  if (n >= 3) return 'CORROBORATED';\n"
    "  if (n >= 1) return 'PARTIAL';\n"
    "  return 'UNVERIFIED';\n"
    "}"
)

add_para(
    "These five are not add-ons; they complete the offline-native premise. 10.1 (mesh) is the "
    "structural one — without it, a total shutdown still blinds the board. 10.2–10.5 make retries, "
    "bandwidth, edits, and trust survive that same blackout.",
    bold=True,
)

add_heading("10.6 Advanced & Differentiating Mechanisms (push for uniqueness)", 2)
add_para(
    "The five improvements above make Nidaa offline-native. To be genuinely novel rather than a "
    "re-assembly of known parts, four mechanisms from adjacent research fields can be adopted — each "
    "is established in its own domain but essentially unused in conflict-zone coordination boards:"
)
add_bullets([
    "Signed append-only feeds (Secure Scuttlebutt / SSB pattern): instead of mutable entries, each "
    "device keeps an append-only, cryptographically signed log of its posts and edits; replication is "
    "gossip of log tails. This gives tamper-evidence and natural conflict detection for free, and "
    "needs no server. SSB itself (patchwork, 3.5k★) proves the pattern works offline — but no "
    "humanitarian board uses it. This is the deepest differentiator.",
    "CRDT field-merge (Automerge / Loro / Yjs, all >5k★): replace the hand-rolled three-way merge in "
    "10.4 with a real CRDT so concurrent offline edits to different fields merge automatically and "
    "deterministically, with no central arbitrator — the mathematical guarantee 10.4 only approximates.",
    "Geo-indistinguishable location (differential privacy, Google/Meta/IBM libs exist): when a post "
    "carries a coordinate, add calibrated noise so the true point stays within ~k meters with "
    "cryptographic guarantee, defeating location-targeting (Threat 12.1) while keeping the post useful "
    "for 'nearest aid' ranking. Turns the safety mitigation into a math property.",
    "Proof-of-personhood for sybil resistance (privacy-preserving, e.g. Anon Aadhaar / ZK): cap "
    "corroborations (10.5) per unique human, not per device/account, with a zero-knowledge proof that "
    "does not reveal identity. Closes the trust-forgery threat without a central ID authority — "
    "critical when no government is trusted.",
    "On-device LLM triage & translation (WebLLM / transformers.js, runs in-browser): a small model, "
    "downloaded once, translates posts between Arabic/English/ Ukrainian/etc. and flags urgent medical "
    "posts offline — no cloud, no data leaving the device. Unique for a crisis board and resilient to "
    "shutdown.",
])
add_para(
    "Why these raise the bar: individually each is known; combined into one offline-first, "
    "Arabic-RTL, trust-decentralized board, they push Nidaa from 'a competent offline app' toward a "
    "genuinely novel coordination primitive. The SSB-style signed-feed + CRDT-merge pair, in "
    "particular, is the technical heart that would make 'first-in-combination' a real claim, not a "
    "marketing one.",
    bold=True,
)

# ===== 11. Competitive Landscape — are we first? =====
add_heading("11. Competitive Landscape — Are We First?", 1)
add_para(
    "An exhaustive scan (GitHub repository search + humanitarian-OSS landscape, July 2026) to test the "
    "'we might be first' question. Important framing: the authors are NOT from Syria or Gaza. Nidaa is "
    "therefore scoped as a GENERALIZABLE offline-first humanitarian board — Syria/Gaza are the first "
    "deployment targets, not the only ones. The 'first' claim must therefore hold across ALL conflict "
    "and crisis regions (Ukraine, Sudan, Myanmar, Haiti, Yemen, Afghanistan, plus natural disasters), "
    "not two. The scan below covers exactly that."
)
add_para("What exists today (the real neighbors, by evidence):", bold=True)
add_bullets([
    "crisis-mesh-messenger (192★, active Oct 2025) — decentralized, infrastructure-free MESSAGING via Bluetooth/WiFi Direct, E2E-encrypted. Solves communication when infra fails. But it is chat, NOT a coordination board: no needs/offers posts, no verification, no map, no Arabic.",
    "HTBox/allReady (885★, but last push 2022) — disaster PREPAREDNESS campaigns for orgs. Not conflict-zone, not Arabic, not offline-first board.",
    "melizeche/ayudapy (118★, active) — Spanish 'help people help people' platform. Peacetime mutual aid, not conflict/offline.",
    "Sahana Eden (409★) — national-scale IMS. Needs institutional backing a local committee lacks.",
    "Kiwix (1,382★) / Rushlight — offline CONTENT (Wikipedia, maps), not live coordination.",
    "HotOSM MapSwipe (170★) / fAIr (137★) — crowdsourced MAPPING, not posting/coordination.",
    "Mutual-aid OSS (factn/resilience-app 85★, rubyforgood/mutual-aid 55★, Crown Heights 26★) — peacetime community organizing, online-only.",
    "Disaster-response AI repos (Gemma/CRISIS-AI etc.) — all 2–13★ hackathon projects, none mature, none Arabic boards.",
])
add_para(
    "Targeted negative evidence (what does NOT exist):", bold=True
)
add_bullets([
    "\"mutual aid / needs board / community board + offline\" → 0 GitHub results. No public OSS project combines those exact concepts.",
    "Arabic-RTL + refugee/aid/offline search → 1 result, and it is a UNHCR refugee-treaty DOC TEMPLATE (1★). No Arabic-RTL humanitarian coordination board exists on GitHub.",
    "Region-specific crisis boards (Ukraine, Haiti, Myanmar, Sudan, Yemen) → none found; Ukraine tech is cybersecurity (uashield 1,077★) and dev tooling, not civilian coordination.",
    "Mesh + coordination (not just messaging) → only 12★ hobby repos (Relief-Mesh, THOR); none matured into boards.",
])
add_para(
    "Verdict: Nidaa is not 'first' in the abstract — offline sync, mesh messaging, aid boards, and "
    "Arabic i18n all exist separately (crisis-mesh-messenger proves mesh-for-crisis is real and "
    "valuable). But the specific COMBINATION it targets — an offline-first, Arabic-RTL community "
    "needs-&-offers BOARD, seeded with real HDX facility data, with mesh sync (Phase 6), decentralized "
    "trust (10.5), and a written threat model (12) — does not appear to exist as a public, deployable "
    "tool in ANY conflict/crisis region. Nearest neighbors are either messaging-only (crisis-mesh-"
    "messenger) or peacetime mutual aid (ayudapy, allReady). So the honest, defensible claim is: "
    "first-in-combination, not first-in-principle — and the gap is wide enough that execution, not "
    "novelty, is the real risk. The multi-region scan (not just Syria/Gaza) strengthens this claim "
    "rather than weakens it: the absence is global, not regional.",
    bold=True,
)

# ===== 12. Threat Model & Do-No-Harm =====
add_heading("12. Threat Model & Do-No-Harm", 1)
add_para(
    "A coordination board in a conflict zone is a dual-use system: it helps civilians find aid, but "
    "the same board can be weaponized — false posts can misdirect or trap, location data can be used "
    "to target, and a 'verified' label can be forged to manufacture trust. Section 10.5 (decentralized "
    "trust) reduces single-point capture but also lowers the bar for bad actors to seed plausible "
    "lies. This section is the deliberate counterweight. It is written before scale, because retrofitting "
    "safety into a trusted tool is far harder than designing for it."
)

add_heading("12.1 Threats", 2)
add_bullets([
    "Misinformation / entrapment: a falsified 'shelter' or 'aid point' post draws civilians to a location. Highest harm, easiest to execute.",
    "Surveillance / targeting: posted coordinates, contact numbers, and movement patterns can be harvested by hostile actors to target aid seekers or volunteers.",
    "Trust forgery: in decentralized trust (10.5), sybil accounts mass-corroborate a fake entry to push it to CORROBORATED/VERIFIED tier.",
    "Denial of service: floods of junk posts bury real needs; mesh forwarding (10.1) can amplify a spam storm across devices.",
    "Server / NGO account compromise: a captured admin account flips verifies en masse or deletes real entries.",
    "Device seizure: a confiscated phone reveals a user's posting history and contacts if stored in the clear.",
])

add_heading("12.2 Mitigations (by design, not afterthought)", 2)
add_bullets([
    "On-device encryption at rest (IndexedDB/OPFS) so a seized device does not leak history (Amplified Access field principle).",
    "No mandatory PII: contact fields are optional; default posts carry no phone/location unless the author chooses. Location shown at city granularity, not street.",
    "Rate-limit + proof-of-work on post creation to blunt spam floods and sybil corroboration.",
    "Corroboration weighting: require distinct device/network fingerprints (not just accounts) for trust tiers; cap one corroboration per device.",
    "Audit log + reversible verifies: every verify/corroboration is logged with actor + time; a compromised admin's actions are visible and revertible.",
    "Mesh forwarding caps (10.1): max hops + duplicate suppression so a spam storm cannot loop forever across the local network.",
    "Clear 'unverified = unconfirmed' UI language so users never treat an unverified post as fact (already in the card design).",
])

add_para(
    "Do-no-harm principle: when a feature trades off safety against convenience (e.g. precise "
    "locations, open posting, lightweight auth), the default MUST favor safety. Nidaa is a civilian "
    "aid tool; it must never become a targeting or misinformation vector. The verification endpoint "
    "(Section 7) and decentralized trust (10.5) are explicitly gated behind this threat model.",
    bold=True,
)

# ===== 13. Conclusion =====
add_heading("13. Conclusion", 1)
add_para(
    "Nidaa is built on a single, defensible principle: the people who need coordination most are the "
    "ones with the least connectivity, so the tool must work without it. By making the device the "
    "source of truth and the network an opportunistic extra, Nidaa stays useful through shutdowns, "
    "damage, and intermittent service — in Syria today, and in Gaza tomorrow."
)
add_para(
    "The path forward is not more features; it is trust, verification, and field use. Those are the "
    "gaps between a working prototype and a tool a community can depend on."
)

# ===== 14. Build Log: Failures, Fixes & Verification Reality =====
# This section exists because a journal that only lists successes is a brochure, not a record.
# The point of Nidaa's journal is to be HONEST about how the tool was actually built — what broke,
# what we wrongly thought was fixed, and what we could and could not verify.
add_heading("14. Build Log — Failures, Fixes & Verification Reality", 1)
add_para(
    "This section is the honest engineering record. It lists what failed, what we believed was fixed "
    "but was not, and what was actually verified vs. assumed. If you only read one part of this "
    "document, read this one — it is the difference between a demo and a dependable tool."
)

add_heading("14.1 — 2026-07-13: Phase 1 — Real HDX Data (what broke)", 2)
add_para("Failures encountered while importing real facility data:")
add_bullets([
    "Renamed data/seed.js to .ts but left JS syntax inside → build failed. Fixed by rewriting as TS.",
    "Imported types via './types' instead of '../lib/types' → module-not-found. Fixed.",
    "Gaza/West Bank (PSE) HDX data ships Gaza + West Bank bundled; we initially tagged ALL PSE as "
    "'gza', which would have wrongly shown West Bank facilities (Jenin, Ramallah) under 'Gaza'.",
    "Point-in-polygon clipping against the Gaza Strip boundary failed silently at first: the importer "
    "fell back to tagging everything 'gza' because the boundary file failed to load.",
    "ROOT CAUSE of that silent failure: used `fs.readFileSync` but imported `fs` as the PROMISES "
    "namespace (`import { promises as fs }`), which has no sync method → threw → caught → fallback. "
    "Fixed by importing a sync `fs` for the boundary read.",
    "Polygon nesting was assumed wrong (`coords = [[ring]]`); actually `coords = [ring]`. The clipping "
    "passed a single coordinate where a ring was expected → garbage → all points fell through.",
    "Earlier SHP→GeoJSON converter had an off-by-8 record-length bug and a miscounted bbox offset "
    "(32 vs 36 bytes) that caused an infinite loop on the 114 KB boundary file. Fixed both offsets.",
])

add_heading("14.2 — 2026-07-13: Phase 2 — Offline Maps (what broke)", 2)
add_para("Failures encountered while adding the Leaflet map:")
add_bullets([
    "Imported 'leaflet' types missing → added @types/leaflet.",
    "Dynamic `import 'leaflet/dist/leaflet.css'` inside TSX broke tsc → moved CSS import to layout.tsx.",
    "L.TileLayer.extend() constructor typing too strict → cast instantiation to any.",
    "THE OFFLINE CACHE WAS DEAD: a cached tile blob set `tile.src` but never called Leaflet's "
    "`finish(null)`, so the cached tile never painted and the object URL leaked. Offline, the map "
    "would have shown blank. Fixed by signalling completion and revoking the URL on load.",
    "Map markers LEAKED: code used `eachLayer` + `instanceof L.Marker`, but the map draws "
    "`L.circleMarker` (a different class), so old markers were never removed — every render piled "
    "on more. Fixed with a tracked L.layerGroup + clearLayers().",
])

add_heading("14.3 — 2026-07-13: The region filter — a bug we shipped, then caught", 2)
add_para(
    "When the region filter was first added, its predicate was "
    "`(e.region && e.region === region) || !e.syncedAt`. Because the API serves every entry with "
    "`syncedAt: null` (the importer never sets a real sync time), the `!e.syncedAt` clause evaluated "
    "TRUE for ALL entries — the filter did nothing (selecting 'Gaza' still showed 10,382 entries). "
    "This passed a casual look but failed the count test. Fixed to filter by `region` for entries "
    "that have one, keeping region-less user posts always visible."
)

add_heading("14.4 — 2026-07-13: Concurrency bug — lost posts under load", 2)
add_para(
    "The JSON store did a read-modify-write with no synchronization. Two concurrent POSTs could each "
    "read the old db.json and overwrite the other's upsert, silently losing an entry. A test firing "
    "15 concurrent posts lost 14. Fixed with an in-process promise-chained write lock. After the fix, "
    "20 concurrent posts → 20 survived, zero lost."
)
add_para(
    "A SECOND, related failure: during concurrency testing on a stale server process, the unlocked "
    "writes corrupted db.json (overlapping partial JSON), which then 500'd every GET with "
    "'Unexpected non-whitespace character after JSON'. This was a test-environment artifact, not a "
    "shipping defect — but it proved the race was real and motivated the lock."
)

add_heading("14.5 — 2026-07-13: Verification honesty — what we could and could NOT prove", 2)
add_para("To avoid overclaiming, here is exactly what was and was not verified:")
add_bullets([
    "VERIFIED: npm run build passes; npm run lint passes (eslint + eslint-config-next installed — "
    "note: lint was NOT wired up for the first several sessions; 'lint clean' was a false positive "
    "until eslint was actually installed).",
    "VERIFIED: runtime API GET/POST round-trip; Arabic (UTF-8) titles persist correctly; known "
    "facilities (e.g. a Gaza City / Daraa hospital) present and searchable by city.",
    "VERIFIED (2026-07-13): region split after clipping — Gaza Strip 1,315 / West Bank 3,397 / "
    "Syria 5,670 (+5 illustrative demo seeds present at that time). CORRECTION: the 5 demo "
    "seeds were later STRIPPED from the live board by fix B (commit 09a8573, 2026-07-14); "
    "the live board now holds 10,382 real HDX facilities and 0 fabricated entries.",
    "VERIFIED: concurrent-write lock (20/20 survived).",
    "NOT VERIFIED IN-BROWSER: the live card/map render could not be visually confirmed in this "
    "environment because the harness's browser sandbox blocks the client-side fetch to the local "
    "API (it intercepts it as a 'sensitive primitive'). The same code rendered correctly in an "
    "earlier check; the build is green; the filter logic is proven against the real API payload. "
    "This is a harness limitation, not a known code defect — but it means the visual layer is "
    "asserted, not eyeballed, here.",
    "NOT YET BUILT (planned, not claimed done): Phase 3 auth + role-gated verify (the open verify "
    "endpoint is still open by design for the demo; Phase 3 closes it), SQLite store, conflict "
    "merge, mesh sync, SSB/CRDT/DP/ZK mechanisms.",
])

add_heading("14.6 — 2026-07-13: What 'done' actually means here", 2)
add_para(
    "As of this writing, Nidaa is a verified prototype: real data, offline-first sync, a working "
    "offline-cacheable map, a region scoping filter, and a concurrency-safe store. It is NOT yet a "
    "field-deployable tool — the verification endpoint is open, there is no authentication, no "
    "mesh sync, and no decentralized trust. Those are the honest remaining gaps, and they are the "
    "subject of Phases 3–7, not hand-waving."
)

add_heading("14.7 — 2026-07-13: Third bug pass — the PWA was not actually installable", 2)
add_para(
    "In a later verification pass we found the app's offline-first PWA claim was partially false. "
    "`public/` was completely empty, yet `sw.js` cached `/manifest.json` and `/icon.svg` and the "
    "root layout set `metadata.manifest` to `/manifest.json`. With those files missing, the service "
    "worker's `caches.addAll([\"/\", \"/manifest.json\", \"/icon.svg\"])` rejected (silently caught), "
    "and there was no manifest for the browser to install from — so the app was NOT installable as a "
    "PWA and had no tab icon, despite the offline app-shell caching working. Fixed by creating a real "
    "`public/manifest.json` (RTL Arabic, green-on-navy theme, SVG icon) and `public/icon.svg` (a "
    "map-pin + 'ن' mark). Verified: all three assets serve HTTP 200 and the manifest parses as valid "
    "JSON, so `addAll` now succeeds and the app is genuinely installable offline."
)
add_para(
    "Same pass also corrected two leftover Gaza-first inconsistencies: the document meta description "
    "still said 'for Syria' (now 'Gaza, the West Bank, and Syria'), and the new-post city placeholder "
    "was hardcoded to 'Aleppo' (now a Gaza example)."
)

add_heading("14.8 — 2026-07-13: Map flickered / disappeared while zooming", 2)
add_para(
    "Reported symptom: on the board map, tiles vanished during a zoom gesture and reappeared after. "
    "Root cause was NOT Leaflet — it was our CUSTOM tile layer (lib/tileCache.ts). Its createTile() had "
    "a broken lifecycle: for a cached tile it called Leaflet's finish() BEFORE the blob image decoded, "
    "and set `tile.onload = () => URL.revokeObjectURL(objUrl)` — so the object URL was revoked while "
    "Leaflet was still using that tile element in the zoom animation. The tile blanked out, then "
    "reloaded on the next frame -> flicker. Custom createTile() overrides are fragile across Leaflet's "
    "zoom animation."
)
add_para(
    "Fix (KISS): deleted the custom IndexedDB tile layer and switched BoardMap to the STANDARD "
    "`L.tileLayer(...).addTo(map)`, which handles zoom correctly. Offline tiles are now cached by the "
    "service worker instead (sw.js: static/cross-origin tile requests are cache-first with write-thru "
    "on network success), and the SW cache was bumped to nidaa-v3 so the old shell/tiles are discarded. "
    "Removed lib/tileCache.ts (dead code). Verified: `makeCachedTileLayer` is absent from the build; "
    "`L.tileLayer` + openstreetmap present; npm run build + lint green."
)
add_para(
    "Honest caveat: the live zoom could not be eyeballed in the build environment (the browser sandbox "
    "blocks the client-side fetch to the local API, so the map does not paint for us here). The flicker "
    "was a code-level lifecycle bug in the custom layer, now removed; the standard layer does not have "
    "it. User-side confirmation (hard-reload + zoom) is the final check."
)

add_heading("14.9 — 2026-07-14: From build mode to learning mode (failures, disagreements, red meetings)", 2)
add_para(
    "This session was almost entirely non-code. We challenged the project's direction and, in doing so, "
    "surfaced failures of our own earlier thinking. Recorded honestly:"
)
add_bullets([
    "FAILURE OF FRAMING (morning): the demo video we built (operational Palantir/NYT style, 92s, Arabic+English "
    "burned subs) was scrapped on user instruction ('scrap the video for gods sake'). Lesson: a video is not "
    "the right artifact for this audience; a document outperforms it for NGO coordinators, technical volunteers, "
    "and OSM contributors who skim, forward, and translate. We replaced it with a Word FAQ/overview package.",
    "DISAGREEMENT / CHALLENGE: the user pushed me to challenge my own 'no fatal flaw' conclusion from the red-team "
    "review. I had accepted the red team's verdict too comfortably. On re-examination I found the verdict was a "
    "property of the DOCUMENT, not a deployed system, and that 'governance has no fatal flaw' was really "
    "'governance is unfalsifiable' — an untested assumption wearing a design. I also caught that Nidaa as built "
    "solves a NARROWER problem (resilient local list) than the pitch (cross-community coordination), because "
    "mesh/federation is not built.",
    "RED MEETING (delegated red-team, 4 personas): a subagent attacked the outreach doc as Gaza coordinator, "
    "diaspora organizer, HOTOSM validator, and NGO program manager. It found real INTERNAL CONTRADICTIONS in our "
    "own docs: (1) 'another verifier can undo it' vs 'at least one designated verifier' — with one verifier "
    "there is no fallback; (2) 'governance model: Implemented' vs trust-bootstrap is an open problem; (3) "
    "self-host/E2E/open-source presented as the targeting-risk fix when NONE are built; (4) verification is "
    "online-only, so trust dies exactly during the shutdowns Nidaa is built for — a contradiction with the "
    "offline-first pitch we had never admitted; (5) geolocated needs board is a targeting hazard, buried in Q34; "
    "(6) stale HDX/OSM facility data shown as real, no date — lethal in active conflict; (7) 'not a partnership' "
    "hid asymmetric risk/labor.",
    "FAILURE WE CORRECTED: I initially told the user the seed data was 'real operational data' (10,387 entries). "
    "On checking the code I found seed.ts entries are EXPLICITLY ILLUSTRATIVE Syria data, not operational. The "
    "FAQ/overview now states this plainly. Another self-correction: the red team under-ranked 'no evidence the "
    "problem is a priority' as weakness #5; I re-ranked it as the ROOT weakness.",
    "ACHIEVED (documents): rewrote the FAQ/overview docx with the red-team fixes — Deployed-vs-Promised controls "
    "table, verification-online-only admission, targeting hazard elevated to PRIMARY ethical risk, ≥2 verifiers "
    "required, governance relabeled, failure ranking led with physical harm. Added NIDAA-REDTEAM-CRITIQUE.md as "
    "evidence.",
])
add_para(
    "The strategic pivot that mattered: we stopped optimizing outreach around SECURING A PILOT and started "
    "optimizing it around DISCOVERING WHETHER THE PROBLEM IS REAL. The goal of the first five conversations is "
    "to FALSIFY Nidaa's assumptions, not validate them."
)

add_heading("14.10 — 2026-07-14: Falsification discipline & outreach readiness achieved", 2)
add_para("Concrete artifacts produced and committed this session (all markdown/docx, no app code changed):")
add_bullets([
    "ASSUMPTION-TRACKER.md — 7 core assumptions (A1–A7) with Evidence For/Against, Confidence, Status; a hard "
    "PIVOT RULE (if A1/A2/A5 strongly falsified by multiple conversations → pivot/narrow/stop); and a mandatory "
    "post-conversation template that forbids 'they liked it' summaries.",
    "LEARNING-PLAN.md — reframed milestones: Day 30 = conversations + populated tracker + A1/A2/A5 read; Day 60 = "
    "patterns + classification; Day 90 = formal evidence-based DECISION (proceed/narrow/pivot/stop), not 'first "
    "pilot'. Success redefined as reducing uncertainty, not obtaining a pilot.",
    "PREDICTIONS.md — baseline priors (A1 60%, A2 55%, A3 65%, A4 50%, A5 60%, A6 45%, A7 40%) written BEFORE "
    "evidence to prevent hindsight bias.",
    "first-touch-messages.md — rewritten as problem-discovery (expertise not commitment; 'Nidaa may be solving "
    "the wrong problem'; invite correction); plus a 23-question interview guide built to break assumptions.",
    "OUTREACH-TARGETS.md + STAKEHOLDER-QUESTIONNAIRE.md + OUTREACH-CONTACTS.md — 4 targets, fill-in drafts (ask "
    "= written brief, not a call), and VERIFIED contact surfaces from live site checks (Sameer: "
    "info@thesameerproject.com; HOTOSM: info@hotosm.org + WNAHub@hotosm.org; HEAL email not found → CONFIRM via "
    "social; Gaza Soup Kitchen is NOT one org → de-prioritized). No email was guessed.",
    "PILOT-READINESS-BACKLOG.md — Track B frozen reference; ranked by outcome-resilience; excludes mesh/federation/"
    "AI/reputation/moderation/cross-community; Verifier Identity named top item but flagged assumption-LIGHT (depends "
    "on A6, not A1/A2/A5). No implementation authorized until evidence arrives.",
])
add_para(
    "STATE AT SESSION END: planning loop closed. The tracker is the source of truth. Progress is measured by "
    "evidence, not output. The next meaningful artifact is assumption-log/<ORG>-<DATE>.md from a real conversation. "
    "Outreach (send Message #1) is the correct action until Conversation #1 exists. Reality now has priority over "
    "planning."
)

# ===== 14.11 — 2026-07-14 (evening): Evidence collection has started =====
add_heading("14.11 — 2026-07-14: Evidence collection has started (first interview logged)", 2)
add_para(
    "This session marks the project's transition from internal planning into EXTERNAL EVIDENCE "
    "collection. Per the posture agreed 2026-07-14, this is evidence-collection mode, not build mode: "
    "no new strategy docs, no new plans, no feature work. Progress is measured by conversations and "
    "what they do to the assumptions."
)
add_bullets([
    "FIRST INTERVIEW LOGGED: Adam Elijilah (resident perspective), 2026-07-14 — "
    "assumption-log/AdamElijilah-2026-07-14.md. One resident conversation produced more usable signal "
    "than many hours of internal planning.",
    "KEY FINDING (sharpens the thesis): incumbent tools (WhatsApp, Telegram, Facebook) WORK when "
    "connectivity exists; their failure mode is connectivity ITSELF, not a feature gap. The precise "
    "framing is NOT 'WhatsApp is insufficient' — it is 'the tools fail exactly when the network drops, "
    "which is when they are needed most.' Aid info is pushed directly to beneficiaries' phones, so an "
    "outage severs the channel entirely (one family lost access to food-distribution info).",
    "WIDER BLAST RADIUS observed: digital payments/transfers also go unavailable during outages — not "
    "just coordination info. The outage cost is broader than coordination alone.",
    "ASSUMPTION MOVEMENT: A3 (offline matters) STRONGLY strengthened; A2 (info access is a bottleneck) "
    "and A5 (existing tools insufficient for some communities) MODERATELY strengthened. NONE weakened, "
    "none falsified. Pivot Rule NOT triggered (A1/A2/A5 not falsified).",
    "OPEN DEPENDENCY surfaced: Nidaa's value during an outage only exists if aid info is POSTED into a "
    "local-first board first. Incumbent push-channels don't require that publishing step — so the "
    "entry/publishing workflow (who posts? adopted by A7) is a real precondition, not a freebie.",
    "OUTREACH: 9 messages sent (Sameer Project, HOTOSM, HOTOSM WNA Hub, Hand4Gaza, Ali AbuAlatta, "
    "Shuruq As'ad, Adam Elijilah, Saed Al Farra); 1 pending (Israa Zumili). See OUTREACH-CONTACTS.md "
    "status log.",
    "NEXT PHASE: pattern detection across conversations, NOT feature development. Log evidence only; "
    "do not treat likes / read-receipts / follows as evidence. Await operations / NGO / journalism / "
    "field-coordination perspectives as the highest-value untested signal.",
])

# ===== 14.12 — 2026-07-14 (night): record-correction pass (honesty fixes + infra files) =====
add_heading("14.12 — 2026-07-14: Record correction — honesty fixes and new infra files", 2)
add_para(
    "Corrective journal entry only. No new features, no new strategy. Restores alignment between "
    "repo state, runtime behavior, and journal claims."
)
add_bullets([
    "A/B/C HONESTY FIXES (commit 09a8573, 2026-07-14): (A) the store no longer fabricates data when "
    "data/db.json is absent — it starts EMPTY and the API returns setupRequired:true; (B) illustrative "
    "demo seeds are stripped from the live board, which now holds 10,382 real HDX facilities and 0 "
    "fabricated entries (see 14.5 correction); (C) the verification audit log now uses its own write "
    "lock, matching the primary store. Verified: lint + build green; empty-board path returns no "
    "fabricated data; 25 concurrent verifies wrote 25 distinct audit lines with zero loss.",
    "NEW OUTREACH/RESEARCH INFRA FILES added to the repo (process only, no product change): "
    "OUTREACH-CONTACTS.md now holds a 34-org 'Grassroots & Community Outreach Candidates' section "
    "(commit 86550f4); RESEARCH-BACKLOG.md captures unanswered questions from conversations "
    "(commit e24ddf1); OUTREACH-LEDGER.md tracks approached vs to-be-approached with a five-state "
    "machine (commits fd00953, c63cb0e). Mones (Humanitarian Ops, MSF) added to the outreach log "
    "(commit d43412d).",
    "OUTREACH COUNT corrected: 9 messages sent, 1 pending (Israa Zumili), 1 interviewed (Adam). See "
    "14.11 correction above.",
    "STATUS: these changes align the journal with runtime behavior. No further journal updates until a "
    "new interview, a tracker change, or another correctness issue. Evidence-collection mode continues.",
])

# ===== 15. Research & Validation History =====
add_heading("15. Research & Validation History", 1)
add_para(
    "This section is the project's evidence memory. It records what we believed before evidence, what "
    "evidence has arrived, and what remains untested. It is synthesized from ASSUMPTION-TRACKER.md, "
    "PREDICTIONS.md, LEARNING-PLAN.md, and the interview log — not reinterpreted."
)
add_heading("15.1 Why the assumptions (A1–A7) were created", 2)
add_para(
    "Nidaa began as a build. The shift to evidence collection produced a set of 7 core assumptions to "
    "be actively falsified, not validated (ASSUMPTION-TRACKER.md). The framing was explicit: the first "
    "conversations are an attempt to FALSIFY the thesis, because falsification saves us from building "
    "the wrong thing. The 7 assumptions:"
)
add_bullets([
    "A1 — Coordination is a priority problem. (baseline belief 60%; Untested as of 2026-07-14)",
    "A2 — Information/matching is a significant bottleneck. (55%; Strengthened)",
    "A3 — Offline capability materially matters. (65%; Strengthened)",
    "A4 — Verification materially matters. (50%; Untested)",
    "A5 — Existing tools are insufficient for at least some communities. (60%; Strengthened)",
    "A6 — Organizations would trust designated verifiers. (45%; Untested)",
    "A7 — Communities would adopt a new workflow if it solved the problem. (40%; Untested)",
])
add_para(
    "PREDICTIONS.md was written BEFORE any conversation and frozen — its priors are not edited after "
    "evidence; reality gets a vote. The baseline was deliberately skeptical on trust (A6 45%, A7 40%) "
    "and uncertain on whether coordination is even the binding problem (A1 60%, Low confidence)."
)
add_heading("15.2 Prediction baselines (written pre-evidence, 2026-07-14)", 2)
add_para(
    "Each assumption carried explicit strengthening / weakening / falsifying conditions. Example: A1 "
    "is falsified only if, across 5+ conversations, coordination is never raised as a priority and is "
    "explicitly subordinate to supply/access/security. A3 is weakened only if disconnection proves "
    "brief/rare and orgs tolerate it with existing tools. These thresholds are the bar evidence must clear."
)
add_heading("15.3 Evidence gathered to date", 2)
add_para(
    "One interview completed: Adam Elijilah (resident perspective), 2026-07-14 — "
    "assumption-log/AdamElijilah-2026-07-14.md. What it established (verbatim from the evidence ledger):"
)
add_bullets([
    "A3 STRENGTHENED (strongly): outages cut access to aid info AND digital payments/transfers; one "
    "family lost access to food-distribution info. The failure mode of incumbent tools IS connectivity loss.",
    "A2 STRENGTHENED (moderately): info access depends on connectivity; when it drops, info access is the binding constraint for receiving aid.",
    "A5 STRENGTHENED (moderately): WhatsApp/Telegram/FB work only when connectivity exists; insufficient precisely during outages, when needed most.",
    "A1, A4, A6, A7 — UNTESTED. No conversation has touched them.",
    "Pivot Rule NOT triggered (A1/A2/A5 not falsified).",
])
add_para(
    "Surprise recorded in the interview: the framing 'WhatsApp is insufficient' is imprecise. The sharper "
    "lesson is that incumbent tools work when connectivity exists and fail exactly when it doesn't — the "
    "failure mode is connectivity itself, not a feature gap. Also: digital payments/transfers fail during "
    "outages too, a wider blast radius than coordination alone."
)
add_heading("15.4 The falsification discipline", 2)
add_para(
    "LEARNING-PLAN.md reframed success: not 'obtain a pilot,' but 'reduce uncertainty enough to make a "
    "high-confidence decision.' The Day 90 goal is a formal evidence-based decision (proceed / narrow / "
    "pivot / stop), not a first pilot. Operating rule: every conversation updates the tracker BEFORE next "
    "steps; no 'they liked it' summaries; the Pivot/Stop conditions are a hard gate. Non-evidence (likes, "
    "read receipts, follows) is explicitly excluded from the record."
)
add_heading("15.5 What lacks evidence / remains hypothesis", 2)
add_bullets([
    "A1 (coordination is a priority) is still Untested — one resident interview did not address whether coordination is the BINDING constraint vs supply/access/security.",
    "A4 (verification matters), A6 (orgs trust designated verifiers), A7 (communities adopt a new workflow) — entirely untested.",
    "Whether Nidaa helps during an outage depends on someone POSTING aid info into it first — an entry/publishing dependency incumbent push-channels don't have (touches A7). This precondition is identified but unvalidated.",
    "All multi-conversation patterns are unknown; only one data point exists.",
])

# ===== 16. Decision History =====
add_heading("16. Decision History", 1)
add_para(
    "Decisions that changed the project's direction, with trigger, date, and reasoning. Synthesized from "
    "the build log (14.9–14.12), LEARNING-PLAN.md, and the red-team critique. No conclusions invented."
)
add_heading("16.1 Pilot-first → problem-validation pivot", 2)
add_para(
    "Trigger: a red-team critique (2026-07-14) and the user's instruction to challenge the project's own "
    "direction. Decision: stop optimizing outreach around SECURING A PILOT; start optimizing around "
    "DISCOVERING WHETHER THE PROBLEM IS REAL. The goal of the first five conversations became falsification, "
    "not validation. Recorded in Build Log 14.9–14.10. This is the single largest directional decision in the project."
)
add_heading("16.2 Red-team conclusions and what we corrected", 2)
add_para(
    "A delegated red-team (4 personas: Gaza coordinator, diaspora organizer, HOTOSM validator, NGO program "
    "manager) produced 8 high-priority fixes (NIDAA-REDTEAM-CRITIQUE.md). The most important, carried into the "
    "FAQ/overview rewrite:"
)
add_bullets([
    "Single-verifier requirement contradicted the 'another verifier can undo it' mitigation — no fallback exists.",
    "Headline privacy controls (self-host, E2E, open source) were presented as present but were NOT built; relabeled as planned.",
    "Verification is online-only, so trust dies exactly during the shutdowns Nidaa targets — an unacknowledged contradiction with the offline-first pitch.",
    "Geolocated needs board is a targeting hazard; elevated to a PRIMARY ethical risk.",
    "Stale HDX/OSM facility data shown as real with no date — lethal in active conflict.",
    "Governance labeled 'Implemented' when it was planned/unvalidated; trust-bootstrap an open problem.",
])
add_para(
    "Self-correction during the review: the user pushed back on accepting the red team's 'no fatal flaw' "
    "verdict too comfortably. Re-examination found that verdict was a property of the DOCUMENT, not a deployed "
    "system, and that 'governance has no fatal flaw' was really 'governance is unfalsifiable.' The red team had "
    "also under-ranked 'no evidence the problem is a priority' as weakness #5; it was re-ranked as the ROOT weakness."
)
add_heading("16.3 Outreach-strategy evolution", 2)
add_bullets([
    "Initial: 4 targets with fill-in drafts (Sameer, HEAL, HOTOSM, Gaza Soup Kitchen) — OUTREACH-TARGETS.md.",
    "Correction: Gaza Soup Kitchen is a distributed concept, not one org — de-prioritized; target a specific kitchen only via referral.",
    "Expansion: a 34-org Grassroots & Community Candidates list was added (2026-07-14) to broaden the pipeline; then frozen as sufficient — no further discovery unless the pipeline dries up.",
    "State machine: every contact assigned exactly one of five states (Queued / Contacted / Replied / Interviewed / Closed) to prevent ambiguity at scale — OUTREACH-LEDGER.md.",
])
add_heading("16.4 Scope and governance corrections", 2)
add_bullets([
    "Scope: Nidaa as built solves a NARROWER problem (resilient local list) than the pitch (cross-community coordination), because mesh/federation is not built. This gap is documented, not closed.",
    "Governance: relabeled from 'Implemented' to 'roles defined in code; operating model planned/unvalidated.'",
    "Honesty fixes A/B/C (commit 09a8573): the store no longer fabricates data; demo seeds stripped; audit log write-locked. Treated as correctness, not features.",
])
add_heading("16.5 Pending decision (the Day-90 gate)", 2)
add_para(
    "No proceed/narrow/pivot/stop decision has been made. The Day 90 milestone (LEARNING-PLAN.md) is a formal "
    "decision point contingent on evidence. As of 2026-07-14 only ONE interview exists; the decision is explicitly "
    "deferred. The Pivot Rule remains armed: if A1/A2/A5 are strongly falsified by multiple conversations, pivot/narrow/stop is mandatory."
)

# ===== 17. Failure History =====
add_heading("17. Failure History", 1)
add_para(
    "What was wrong, mistaken, or corrected. Preserved so future readers see the project's real trajectory, "
    "including its errors. Synthesized from the build log and red-team critique. New mistakes, when found, belong here."
)
add_heading("17.1 Wrong assumptions and incorrect reasoning", 2)
add_bullets([
    "Narrower problem than the pitch: Nidaa was presented as cross-community coordination, but mesh/federation is not built — it is a resilient local list. Caught during the red-team self-review (14.9).",
    "Misleading self-statement: an earlier claim that seed data was 'real operational data' (10,387 entries) was FALSE — seed.ts entries are explicitly ILLUSTRATIVE, not operational. Corrected in the FAQ/overview and later by fix B (14.5, 14.12).",
    "Over-accepted red-team verdict: initially accepted 'no fatal flaw' too comfortably; the verdict applied to the document, not a deployed system (14.9).",
    "Imprecise thesis: 'WhatsApp is insufficient' understates the real failure mode (connectivity itself). Sharpened after the Adam interview (15.3).",
])
add_heading("17.2 Misleading documentation (corrected)", 2)
add_bullets([
    "Status box said 'Governance model: Implemented' and 'Verification workflow: Implemented' when both were planned/unvalidated (red-team fix #7). Corrected in the FAQ rewrite.",
    "Privacy mitigations (self-host, E2E, open source) presented as the answer to targeting when none were built (red-team fix #2). Relabeled as planned.",
    "Demo seeds implied as part of the real board; stripped by fix B so the board holds 0 fabricated entries (14.5, 14.12).",
])
add_heading("17.3 Technical mistakes (build era, 2026-07-13)", 2)
add_para(
    "Captured in Build Log 14.1–14.8. Representative examples (all since fixed): renamed seed.js to .ts but left "
    "JS syntax inside → build failed; imported './types' instead of '../lib/types' → module-not-found; point-in-polygon "
    "clipping fell back to tagging all PSE as 'gza' (wrong region split); the offline tile cache signaled completion "
    "incorrectly so cached tiles never painted; map markers leaked via eachLayer + instanceof on the wrong class; the "
    "region filter predicate '!e.syncedAt' evaluated true for ALL entries so the filter did nothing; the JSON store had "
    "an unlocked read-modify-write that lost posts under concurrency (15 concurrent → 14 lost; fixed with a write lock); "
    "the PWA was not actually installable (empty public/, missing manifest/icon). Lint was also reported 'clean' for "
    "several sessions before eslint was actually installed — a false positive."
)
add_heading("17.4 Scope gaps and unresolved concerns", 2)
add_bullets([
    "Verification requires connectivity, so the 'verified' signal dies during the shutdowns Nidaa targets (red-team fix #3). Not yet solved.",
    "Single-verifier has no fallback; 'another verifier can undo it' is impossible with one (red-team fix #1).",
    "Geolocated needs board is a targeting hazard with weak mitigation (red-team fix #4).",
    "No security review, no GDPR/DPIA assessment, no formal do-no-harm policy, data ownership undefined (red-team fix #6).",
    "Silent last-write-wins data loss during conflicts; no field-evidence baseline (red-team fix #8).",
    "Stale facility data shown without provenance/date (red-team fix #5) — import now records source but a refresh procedure is still open.",
])
add_heading("17.5 Failed or closed outreach", 2)
add_para(
    "As of 2026-07-14: ZERO outreach attempts have failed, been declined, or been closed. 9 messages sent, 1 pending "
    "(Israa Zumili, Queued), 1 interviewed (Adam), 0 replied-without-interview, 0 closed. This is recorded explicitly so "
    "the absence is not mistaken for an omission. The pipeline is healthy; the limiting factor is responses, not candidates."
)

# ===== 18. Discussion & Open Questions =====
add_heading("18. Discussion & Open Questions", 1)
add_para(
    "Significant debates and what remains unknown. The journal must make clear what has evidence, what lacks it, and "
    "what is still hypothesis. Uncertainty is not resolved here."
)
add_heading("18.1 Significant debates", 2)
add_bullets([
    "'Scrap the video' (morning, 2026-07-14): a polished demo video was built then discarded on user instruction; a Word FAQ/overview package replaced it as the right artifact for NGO/OSM/diaspora readers. Debate resolved by user call.",
    "'Challenge your own no-fatal-flaw conclusion': the user pushed for adversarial self-review rather than comfort with the red team's verdict. Reshaped the critique's ranking.",
    "'Stop improving the process' (evening, 2026-07-14): after the outreach/ledger/backlog infra matured, the user ruled that further process improvements were diminishing returns; future repo changes gated on a conversation, a repeated operational pain, or a product-correctness issue. The outreach/evidence/tracking systems were declared sufficient.",
])
add_heading("18.2 What we still do not know", 2)
add_para(
    "Open questions are tracked in RESEARCH-BACKLOG.md (seeded from the project framing and from Conversation #1). "
    "They are NOT answered here. Representative open questions:"
)
add_bullets([
    "How do humanitarian orgs currently verify information, and how are beneficiary lists managed during connectivity disruptions?",
    "What trust models are used in community-led aid networks? (bears on A4/A6)",
    "How do field teams coordinate when communications fail? (bears on A3/A5)",
    "What existing offline-first humanitarian tools exist, and do they address the binding constraint?",
    "From Adam's interview: what is the actual outage→info-loss sequence; what do people use when digital payments fail; are there existing offline re-sharing practices; how do resident→humanitarian referral pathways form?",
])
add_heading("18.3 Evidence status summary", 2)
add_para(
    "HAS EVIDENCE: A2, A3, A5 strengthened by one resident interview. LACKS EVIDENCE: A1, A4, A6, A7 untested. "
    "HYPOTHESIS: that coordination/verification/offline are the binding problems — partially supported by one data "
    "point, not established. The next meaningful evidence is Conversation #2 (an operations / NGO / journalism / "
    "field-coordination perspective), which is the highest-value untested signal. Until then, the project intentionally "
    "stays quiet and collects evidence rather than building."
)

add_heading("14.11 — 2026-07-15: Credibility fixes + validation-readiness (pre-pilot, not features)", 2)
add_para(
    "The review of prior sessions concluded the next bottleneck is EXTERNAL VALIDATION, not features. Two credibility "
    "risks were closed before any org sees the board, and the interview instruments were sharpened to test the most "
    "dangerous adoption assumption."
)
add_bullets([
    "CREDIBILITY FIX (import-hdx.mjs + types.ts): imported HDX/HOT OSM facility data no longer carries "
    "verified:true. It is provenance-bearing reference data — tagged source + sourceDate, shown as a purple "
    "'Source: HDX/… · <date>' badge (list + map popup), never a '✓ verified' tick. Header/notice/footer copy "
    "changed from '(verified)' to '(reference data, not human-verified)'. This removes the red-team-flagged "
    "false-verification hazard.",
    "SAFE DEFAULT PRECISION (page.tsx): user posts now default to precision:'neighborhood' (city/area shown, exact "
    "coordinates withheld) to reduce targeting risk in active conflict; imported dataset features stay 'exact'. The "
    "red-team's primary ethical risk (geolocated needs board as targeting list) is now mitigated by default.",
    "INTERVIEW GUIDE SHARPENED (STAKEHOLDER-QUESTIONNAIRE.md + first-touch-messages.md): added the PUBLISHING-STEP "
    "block — who actually posts coordination info today, how it flows from knower to needer, and whether a new tool "
    "gets used only if someone remembers to post into it first. This directly tests A7/A3 and the thesis that 'Nidaa "
    "only helps during an outage IF info was posted into it first.'",
    "NEXT-WAVE OUTREACH DRAFTS (first-touch-messages.md): Messages 5–7 for White Helmets (Syria), Sudan Relief Fund, "
    "and Emergency Response Rooms (Sudan) — extending coverage to A1/A4/A6 and a SECOND conflict region to test "
    "generalizability, not just Gaza. ERRs is [MANUAL] (confirm channel before send).",
])
add_para(
    "Scope discipline maintained: no mesh/CRDT/AI/federation work; no new feature logic; all changes are honesty + "
    "safety + evidence-collection. Build green (npm run build exit 0); importer confirmed to leave verified=false."
)
add_para(
    "Open decision recorded: the implementation gate ('no code until Conversation #1 updates the tracker') is now "
    "satisfied, and these credibility fixes are UNBLOCKED. The binding constraint remains replies (7 sent, 0 replied) "
    "— the next artifact is Conversation #2, not more code."
)

add_heading("14.12 — 2026-07-15 (morning): Validation-readiness sprint (no replies yet, keep moving)", 2)
add_para(
    "With no replies landed, the morning was spent on work that does not depend on external responses: broadening the "
    "contact pipeline, making knowledge gaps explicit, prepping for replies, and studying existing coordination tools."
)
add_bullets([
    "NEW CONTACTS: 9 additional verified channels added to OUTREACH-LEDGER.md (#12–#20): REACH (MEAL, verified email), "
    "iMMAP (IM/MEAL, verified email), Ground Truth Solutions (accountability, verified email), Bonyan (Syria local, "
    "MANUAL), Tech for Palestine, MAP, Islamic Relief USA, We Are Not Numbers, ACAPS. Unreachable orgs (Molham, "
    "Darfur Women, Sudan Emergency, Voices of Return, SAPAM — HTTP 000/404) were deliberately NOT added.",
    "ASSUMPTION TABLE (pilot/ASSUMPTION-NEXT-TESTER.md): A1–A7 each with Evidence / Confidence / Next tester. Makes "
    "gaps obvious — A1/A4/A6/A7 still Untested; A2/A3/A5 Medium (1 resident).",
    "TOOL RESEARCH (pilot/COORDINATION-TOOLS-RESEARCH.md): live-fetched Ushahidi, HOTOSM, ReliefWeb, HDX. Key finding "
    "on the 'who posts?' question: mature tools ALWAYS have an owner (deployment owner / validator / OCHA editor / org "
    "data focal point). Beneficiaries rarely post; orgs and coordinators do. Supports designing Nidaa for the "
    "coordinator/info-officer, not the lone beneficiary.",
    "REPLY PREP (pilot/INTERVIEW-REPLY-PREP.md): opening, 6 follow-up questions, note template, assumption-mapping, and "
    "the mandatory post-call pipeline (log → tracker → ledger → journal → commit) — so any reply today converts to "
    "evidence fast.",
    "NEXT-WAVE MESSAGES BATCH 2 (first-touch-messages.md #8–#11): drafts for REACH, iMMAP, Ground Truth, and the "
    "Gaza/Sudan community voices, all verified channels.",
])
add_para(
    "Morning success metric met: 9 new verified contacts, assumptions table live, interview guide + reply prep ready, "
    "and a concrete read on existing workflows. No code changed; no new feature logic. Repo stays evidence-collection."
)

add_heading("14.13 — 2026-07-15 (midday): Thesis shift — A7 splits into A7a/A7b", 2)
add_para(
    "The morning's tool research produced a real thesis refinement, not just more pipeline. The dangerous assumption was "
    "never 'will beneficiaries post?' — it is 'does the partner org already have someone who OWNS information flow?' Every "
    "mature system has an owner: Ushahidi→deployment owner, HOTOSM→validators, ReliefWeb→OCHA editors, HDX→org data managers."
)
add_bullets([
    "A7 split into A7a (coordinators/info officers will publish — now LESS risky; we fit an existing role) and A7b "
    "(beneficiaries open/use the info — still largely untested). Recorded in ASSUMPTION-NEXT-TESTER.md.",
    "PILOT DESIGN implication: target the coordinator/info-officer as primary user; build a bridge (coordinator forwards "
    "from WhatsApp) for communities without a formal focal point. Do NOT design for the lone beneficiary posting.",
    "AFTERNOON-SEND-PACKAGE.md created: one fire-ready list (Tier 0 follow-ups → Batch 1 → Batch 2, verified channels, "
    "reply protocol, thesis-to-test). Solves the 'which draft / which channel / what order' retrieval problem for sending.",
    "Discipline note: infrastructure is now sufficient (ledger, assumption table, tool research, reply prep, journal, "
    "tracker, learning plan, pilot readiness, research backlog). No further framework work authorized — next gain is a reply.",
])
add_para(
    "State: waiting. Highest-value event is no longer a commit — it is a reply from ops/coordination/PMER/IM/community-"
    "response. Afternoon objective: send Batch 1 + Batch 2 and secure Conversation #2."
)

add_heading("14.14 — 2026-07-15 (strategy): coordinator-led hypothesis + failure-mode register", 2)
add_para(
    "Strategic reassessment after the morning sprint and red-team work. New WORKING HYPOTHESIS (explicitly not a "
    "conclusion): Nidaa may be a coordinator-led information-management/coordination tool, not a civilian-first "
    "platform. Recorded in ASSUMPTION-TRACKER.md as a falsifiable hypothesis, with a Potential Failure Modes table "
    "(F1 coordinator overload, F2 coordination not a priority, F3 orgs reject decentralized info, F4 civilians don't "
    "access/trust, F5 staged workflow mismatch) — each with evidence/confidence/next-tester/strengthen/falsify."
)
add_bullets([
    "INTERVIEW OBJECTIVE SHIFT: future conversations exist to FALSIFY, not validate. Behavioral evidence over "
    "opinions — prefer 'tell me about the last time…', 'what actually happened?', 'who did you contact?', 'how did "
    "information move?', 'what did you do when the plan failed?'. Avoid 'would you use Nidaa?' / 'do you like this?'.",
    "HIGHEST-PRIORITY UNKNOWNS recorded: do coordinators maintain structured info; is coordination a top pain vs "
    "fuel/access/security; who owns info in the field; what happens in worst breakdowns; how orgs verify/trust.",
    "GUIDANCE (frozen scope): freeze ALL new architecture discussion (mesh, LoRa, CRDTs, civilian portals) until "
    "organization-level evidence arrives. Bottleneck is lack of field evidence, not technical uncertainty. Success "
    "metric unchanged: org-level evidence strengthening or weakening A1, A6, A7.",
])
add_para(
    "This is the last planned strategy artifact. With the hypothesis, failure modes, interview shift, and unknowns "
    "now canonical, the project is fully in evidence-collection mode and the repository is frozen for features."
)

add_heading("14.15 — 2026-07-15 (midday): 3 targeted lead profiles added to outreach doc + ledger", 2)
add_para(
    "Added a 'Targeted Lead Profiles' section to Nidaa-Outreach-Targets.docx and a matching 'Targeted Lead Profiles' "
    "table to OUTREACH-LEDGER.md (L1–L3). Each profile maps to specific assumptions and failure modes: L1 Field "
    "Ops Coordinator (MSF/IRC, Gaza) tests A1/A3/A5 + F1/F2; L2 Cross-Border Ops Manager (DRC/Goal, Gaziantep) tests "
    "A4/A6 + F3/F5; L3 IMO Cash&Voucher (CCD/WFP CWG) tests A1/A6 + F1/F3."
)
add_bullets([
    "Channels VERIFIED LIVE this session: IRC donorservices@rescue.org (200); ReliefWeb/humanitarianresponse.info "
    "submit@reliefweb.int (200, legitimate cluster/WG route to CWG/ICCG). MSF contact 403 (bot-blocked); DRC 200 "
    "no email scraped; GOAL 404 (structure changed); Cash Learning 000 (excluded pending human verification).",
    "HONESTY BAR: named individuals (former MSF/IRC field coordinators etc.) are NOT fabricated — exact LinkedIn "
    "search strings provided and flagged MANUAL. No guessed personal emails.",
    "Status for all three: Not Sent. Org channels verified where reachable; named individuals + unreachable orgs "
    "require manual confirmation before outreach.",
])
add_para(
    "These are the highest-value org-level leads yet because they hit the coordinator-led hypothesis directly and "
    "can falsify F1/F2/F3/F5 with behavioral evidence. Still waiting on the first reply to run the evidence pipeline."
)

# ===== 14.16 — 2026-07-15 (evening): evidence-ops infrastructure + new reply =====
add_heading("14.16 — 2026-07-15: Evidence-operations infrastructure + new reply (Abdelrahman Saleh)", 2)
add_para(
    "Documentation + record-keeping only. No feature, architecture, mesh, LoRa, UI, or solution work. "
    "Per the frozen-scope discipline, this session is evidence-collection support while real-world replies land."
)
add_bullets([
    "NEW REPLY LOGGED (not yet an interview): Abdelrahman Saleh — founder of the Wa7at Initiative, a "
    "Gaza-based software engineer. Replied on LinkedIn and identified internet + electricity scarcity as "
    "major challenges. Assumptions touched: A2, A3, A5, A7. Structured follow-up questionnaire sent; "
    "detailed evidence pending. Added to OUTREACH-LEDGER.md (#11, state=Replied) and OUTREACH-CONTACTS.md "
    "status log (Replied count 0 -> 1).",
    "NEW FILE pilot/EVIDENCE-INTAKE-TEMPLATE.md — single intake form for every piece of evidence: contact, "
    "role, date, exact quote/story, information origin/path/consumer, bottleneck, workaround, assumptions "
    "affected, evidence strength, strengthens/weakens/falsifies. Enforces the no-impressions honesty bar.",
    "NEW FILE pilot/CONTACT-STATUS.md — live status board for 7 tracked contacts (Omar Ghazal, Shaima Taha, "
    "Abdelrahman Saleh, Ibrahim Haboush, Naji Al Jafarawi, Fatma Raed, iMMAP) with the five columns "
    "Contacted / Replied / Evidence received / Follow-up required / Referral received.",
    "NEW FILE pilot/CURRENT-EVIDENCE-STATE.md — one-pager: tested assumptions (A2/A3/A5 Strengthened), "
    "untested (A1/A4/A6/A7), evidence gathered so far, evidence still required, and the top-3 highest-value "
    "pending conversations (Abdelrahman Saleh; Ali AbuAlatta/Sameer/HOTOSM; iMMAP + Adam referral).",
    "EVIDENCE COUNT as of this entry: 1 full interview (Adam E.) + 1 partial reply (Abdelrahman S.) = 2 "
    "evidence items. OUTREACH COUNT: Contacted 7, Queued 1, Interviewed 1, Replied 1, Closed 1 (11 total). "
    "No assumption weakened or falsified; Pivot Rule not triggered.",
])
add_para(
    "Next action remains evidence collection: convert Abdelrahman's pending detailed response into a logged "
    "entry via the intake template, and pursue the coordinator-level conversations that can move A1/A6/A7."
)

# ===== 14.17 — 2026-07-15 (evening): strategic pivot to Evidence Pyramid (Tier 1 public evidence) =====
add_heading("14.17 — 2026-07-15: Strategic pivot — Evidence Pyramid (public evidence + interviews)", 2)
add_para(
    "Strategy correction, no feature work. The posture moved from 'interviews are the only evidence' to a "
    "hybrid Evidence Pyramid: Tier 1 public evidence (~60% effort), Tier 2 field interviews (~40%, in parallel, "
    "job = triangulation not teaching), Tier 3 solution evidence deferred until A1/A4/A6 survive both tiers."
)
add_bullets([
    "NEW FILE pilot/PUBLIC-EVIDENCE-REVIEW.md — the Tier-1 deliverable. Assumption-by-assumption evidence "
    "matrix (A1 coordination, A4 verification, A3 comms disruptions) built ONLY from REAL, fetchable sources. "
    "9 HDX dataset URLs retrieved and URL-verified live (HTTP 200) this session; ~21 more named targets queued "
    "(OCHA evals, Internews info-integrity, Access Now #KeepItOn, ReliefWeb reports). No citation invented.",
    "HONESTY BAR HELD: every matrix row carries [RETRIEVED] (fetched + URL checked) or [TO RETRIEVE] (named "
    "real source not yet pulled). The doc explicitly states a row with no source is a GAP, not evidence. This "
    "matches the project's no-fabricated-evidence rule.",
    "CONVERGENCE RULE documented: one report weak, one interview weak, three independent sources agreeing = "
    "strong. Interview questions are now DERIVED from public findings ('does this match your reality?'), not "
    "open-ended. This prevents both failure modes: outreach before learning what is known, and reading without "
    "talking to real people.",
    "CURRENT-EVIDENCE-STATE.md updated with the new posture block. A3 remains best-supported (public + Adam + "
    "Abdelrahman converge); A1 weakest (sources show coordination EXISTS as infrastructure, not that it is a "
    "PAIN — needs eval text + field confirmation).",
    "Tier 1 seed established: 9 verified HDX URLs. Target 20-30 sources before Tier 1 declared sufficient. Tier "
    "3 explicitly gated behind A1/A4/A6 survival. No architecture/mesh/UI/solution work performed.",
])
add_para(
    "Next week focus per user direction: EVIDENCE CONVERGENCE, not product design."
)

# ===== 14.18 — 2026-07-15 (evening): A1 evidence campaign + convergence scoreboard =====
add_heading("14.18 — 2026-07-15: A1 evidence campaign + Evidence Convergence Scoreboard", 2)
add_para(
    "Documentation only. No architecture, no solution design, no feature discussion. The milestone is no "
    "longer a prototype — it is 'does A1 survive evidence convergence?'"
)
add_bullets([
    "NEW FILE pilot/A1-EVIDENCE-CAMPAIGN.md — the focused case FOR or AGAINST A1. 8 sections: evidence "
    "supporting (4 structural HDX/OCHA sources), evidence challenging (Adam did not raise coordination as a "
    "pain; mature infra may mean it is solved), gaps, highest-value public sources to retrieve (OCHA evals, "
    "ACAPS/Ground Truth, after-action reviews, ALNAP, academic), highest-value contacts (AbuAlatta, Sameer, "
    "HOTOSM, White Helmets, ERRs, REACH/iMMAP, L1 leads, Adam referral), what strengthens/weakens/falsifies "
    "A1. Per Pivot Rule, A1 falsified by 3+ coordinator interviews + 2+ evals => pivot/narrow/stop.",
    "NEW FILE pilot/EVIDENCE-CONVERGENCE-SCOREBOARD.md — per-assumption counts (Public / First-Hand / "
    "Interview) + Convergence Status. Definitions: Public = structural datasets; First-Hand = published "
    "primary-field accounts; Interview = direct talks/replies. As of this entry: A3 = Emerging Convergence "
    "(only one); A1 = Weak (4 structural, 0 pain evidence); A2/A5 = Weak (single interview each); A4 = Weak "
    "conceptual; A6/A7 = Untested.",
    "CURRENT-EVIDENCE-STATE.md extended with the live scoreboard table + the A1-risk callout. A3 remains the "
    "strongest signal; A1 is explicitly the largest unresolved risk and the gating milestone.",
    "HONESTY BAR HELD: all counts are real (4 A1 public sources = the HDX/OCHA URLs retrieved 14.17; A3 "
    "interview=2 = Adam + Abdelrahman). No assumed counts, no invented sources. Falsification discipline "
    "preserved: A1 is the assumption most likely to kill the thesis, so it gets the focused campaign.",
])
add_para(
    "Next action: drive A1 from Weak to Emerging — retrieve first-hand accounts and secure one coordinator "
    "interview. No design work until A1 resolves."
)

# ===== 14.19 — 2026-07-15 (evening): A1 Resolution Campaign — first-hand AAR retrieved =====
add_heading("14.19 — 2026-07-15: A1 Resolution Campaign — first-hand after-action review retrieved", 2)
add_para(
    "Single A1-focused mission. No new contacts, no design, no architecture, no pilots. Goal: move A1 from "
    "Weak to Emerging Convergence or Falsified via one real after-action review."
)
add_bullets([
    "RETRIEVED ONE REAL AAR: 'Afghanistan — After Action Review: Eastern Region Earthquake Response (31 Aug "
    "2025, Nangarhar/Kunar)', by the Inter-Cluster Coordination Team (ICCT), published by UN OCHA on ReliefWeb "
    "(21 May 2026). URL verified HTTP 200. Method: document review, operational data, KIIs, focus groups, field "
    "consultations with actors + de facto authorities + affected communities.",
    "KEY FINDING (first-hand, primary-field): the AAR's single 'most significant early constraint was NOT a lack "
    "of humanitarian capacity, but limited operational readiness of first 72-hour activation and sequencing "
    "arrangements... contributing to avoidable delays.' A second failure: early Damage Impact Model figures were "
    "constrained by outdated shelter typology + inaccurate intensity data, limiting confidence in damage "
    "estimates. Both are coordination / information-management failures, not capacity failures.",
    "A1 IMPACT: this is direct support for A1 ('when operations fail, coordination is frequently a binding "
    "constraint'). Coordination was the NAMED binding early constraint. It does NOT falsify A1; it strengthens it. "
    "One source only -> not yet convergence (project rule: 3 independent sources = strong). A1 moves Weak -> "
    "Emerging (not yet crossed).",
    "NEW FILE pilot/A1-AFTER-ACTION-ANALYSIS.md — full structure: context, operational failure, root cause, "
    "was coordination involved (yes), primary/secondary/irrelevant (PRIMARY for early constraint), supporting "
    "quotes, impact on A1, comparison vs Adam interview + Abdelrahman reply + existing public evidence. Verdict: "
    "A1 Strengthened (single source; needs 1-2 more independent accounts for convergence).",
    "UPDATED: A1-EVIDENCE-CAMPAIGN.md (added AAR to supporting evidence; status Weak -> Emerging) and "
    "EVIDENCE-CONVERGENCE-SCOREBOARD.md (A1 First-Hand 0 -> 1; status Weak -> 'Weak to Emerging'). Adam's "
    "resident view is complementary (info-to-people layer), not contradictory (coordinator layer).",
    "DISCIPLINE HELD: one real source; no invented convergence; no design work triggered. Gate remains: no "
    "architecture/solution work until A1 reaches Emerging Convergence or is Falsified. Next: 1-2 more independent "
    "coordinator-level accounts or AARs.",
])
add_para(
    "Verdict recorded in A1-AFTER-ACTION-ANALYSIS.md: Based on current evidence, A1 is Strengthened."
)

# ===== 14.20 — 2026-07-15 (evening): A1 Contradiction Hunt — falsification discipline =====
add_heading("14.20 — 2026-07-15: A1 Contradiction Hunt — deliberate attempt to weaken A1", 2)
add_para(
    "Disciplined falsification step. Supportive sourcing paused; the task was to find the strongest credible "
    "case where operations failed primarily due to physical/political/security constraints, not coordination."
)
add_bullets([
    "RETRIEVED ONE REAL CONTRADICTORY SOURCE: IFRC Real-Time Evaluation (RTE) of the Myanmar Earthquake "
    "Response (28 Mar 2025, central Myanmar). URL verified HTTP 200. Formal RTE covering first 3 months; "
    "response led by Myanmar Red Cross (MRCS) + IFRC — coordination was present and reinforced, not absent.",
    "WHY IT BOUNDS A1: the RTE names the binding constraints as 'longstanding limits on humanitarian access "
    "and international presence,' 'widespread infrastructure damage,' 'continued aftershocks,' and 'political "
    "constraints.' 'Coordination' appears only in OCHA/ReliefWeb footer boilerplate, never as a bottleneck. "
    "You cannot coordinate past a restricted operating environment or destroyed infrastructure.",
    "NEW FILE pilot/A1-CONTRADICTORY-EVIDENCE.md — context / failure / root cause / why coordination was not "
    "decisive / impact on A1. Verdict: does NOT falsify A1 (A1 says 'frequently', not 'always') but correctly "
    "BOUNDS it — coordination is binding where capacity+access exist but mobilization/sequencing fails "
    "(Afghanistan); secondary where access/infra/security are themselves the wall (Myanmar).",
    "UPDATED A1-EVIDENCE-CAMPAIGN.md with a dedicated 'Evidence AGAINST A1' (2b) section; scoreboard A1 note "
    "records the contradiction hunt. First-hand accounts now 2 (1 supporting AAR, 1 contradictory RTE). "
    "Supportive sourcing explicitly paused until the contradiction is absorbed.",
    "CONFIDENCE OUTCOME: the strongest evidence against A1 still leaves A1 standing -> confidence INCREASES "
    "(but A1 is now scoped, not sweeping). This is the project's first deliberate self-falsification attempt "
    "and it did its job: it prevented over-generalizing from a single supporting AAR.",
    "NO DESIGN WORK. Gate unchanged: no architecture/solution/pilot until A1 reaches convergence or is "
    "Falsified. Next: one more independent account on EACH side (esp. a coordinator/ops-lead interview ranking "
    "coordination vs fuel/access/security) before declaring convergence.",
])
add_para(
    "A1 verdict after contradiction hunt: Strengthened and bounded (not universal). Project survives its own "
    "attempt to disprove itself."
)

# ===== 14.21 — 2026-07-15 (evening): A1 paused; A6 Resolution Campaign begins =====
add_heading("14.21 — 2026-07-15: A1 paused at bounded conclusion; A6 Resolution Campaign begins", 2)
add_para(
    "Strategic re-prioritization. A1 reached diminishing returns and is preserved at its bounded conclusion. "
    "The new highest existential risk is A6: even if A1 survives, Nidaa dies if coordinators will not maintain "
    "structured information. No Nidaa/features/UI/architecture discussion."
)
add_bullets([
    "NEW FILE pilot/A6-EVIDENCE-CAMPAIGN.md — the question: 'Will coordinators actually maintain structured "
    "information?' Includes the 5 behavioral questions (who enters / how often / under what conditions / what is "
    "worth the effort / what is abandoned under pressure) and the 'what would have to be true to willingly "
    "maintain structured info during a crisis' statement.",
    "A6 EVIDENCE MATRIX built from REAL sources (URLs HTTP 200): SARC IM Department (dedicated owner across 14 "
    "branches, ~40 trained analysts, MEAL unit — FOR, requires capacity building); MapAction automated pipeline "
    "(FOR with caveat: automation exists BECAUSE manual collection is the burden); Ushahidi->HDX (FOR mixed: "
    "needs human cleaning/triage); WFP edge capture (FOR: lightweight laminated-map marking works); HOTOSM "
    "validation + ReliefWeb editorial (FOR: designated human roles). All six point the same way.",
    "A6 READING: structured info maintenance is a REAL, observed behavior — but survives only with a designated "
    "owner + low burden (often automation). The live risk is burden + ownership, not willingness. A6 is bounded "
    "exactly like A1: holds where owner+low burden exist; fails where maintenance is nobody's job under surge.",
    "GAPS HONESTLY NOTED: the grassroots/small-coordinator case (the actual Nidaa user) is NOT yet evidenced — "
    "SARC/WFP are large orgs with departments. And the A6 contradiction hunt (maintenance collapse under peak "
    "load) has NOT been run. Both required before convergence.",
    "UPDATED: EVIDENCE-CONVERGENCE-SCOREBOARD.md (A6 row relabeled to 'coordinators maintain structured info', "
    "status Early Support; reading notes name A6 the top campaign, A1 paused) and CURRENT-EVIDENCE-STATE.md "
    "(added Current Risk Ranking per user's table; refreshed scoreboard). A1 conclusion preserved, not discarded.",
    "NO DESIGN WORK. Gate unchanged for A1; A6 now the active gate. Next: run the A6 contradiction hunt and seek "
    "one small-coordinator / grassroots IM case or interview (the A1 sec.5 contacts still apply).",
])
add_para(
    "A6 verdict (early): Strengthened (early) — real IM behavior observed, but conditioned on owner + low burden; "
    "grassroots case + contradiction hunt pending."
)

# ===== 14.22 — 2026-07-15 (evening): A6 Grassroots Hunt — corrected the population =====
add_heading("14.22 — 2026-07-15: A6 Grassroots Hunt — the formal-org evidence answered the wrong population", 2)
add_para(
    "Direction correction (user-pushed). The first A6 pass proved formal orgs maintain structure — never in "
    "doubt. The live question is whether the SMALL/INFORMAL coordinator Nidaa depends on does. No design/"
    "architecture/solution work."
)
add_bullets([
    "NEW FILE pilot/A6-GRASSROOTS-HUNT.md — hunts the Nidaa-target population only (mutual-aid networks, "
    "neighborhood committees, camp committees, local volunteers, ERRs, White Helmets local, community "
    "organizers). Applies the 6 behavioral questions (who keeps info / how / what written / what verbal / "
    "when breaks / what abandoned first).",
    "REAL GRASSROOTS SOURCES (HTTP 200): ACAPS 'Khartoum State Emergency Response Rooms' (16 Oct 2025) and "
    "ACAPS/SSHAP 'Mutual aid lessons — ERRs Sudan' (case study Jun-Aug 2024). ERRs = ~360 volunteer networks "
    "across 7 states; Khartoum ERRs across 54 neighborhoods; many from neighborhood resistance committees. "
    "The 'room' = the ONLINE GROUP CHATS where they were conceived — coordination is chat/verbal-based, not a "
    "structured database. The brief calls for 'ongoing documentation, analysis and learning' as a RECOMMENDED "
    "GAP, not an observed practice.",
    "GRASSROOTS ANSWER: no dedicated owner; info in chat; beneficiary/distribution data lives in chat not a "
    "register; most coordination verbal; breaks down under surge/insecurity/CONNECTIVITY LOSS (chat-dependent "
    "= inherits A3); documentation abandoned first under load.",
    "SUCCESS CRITERION: NO successful grassroots structured-maintenance example yet found; collapse-under-"
    "pressure partially evidenced (chat-dependence + documentation-as-gap). So for the Nidaa user, A6 is "
    "UNPROVEN, LEANING WEAK — the opposite of the formal-org 'early support.' The formal evidence does NOT "
    "transfer to this population.",
    "IMPLICATION (behavioral, not willingness): structured maintenance by small coordinators is plausible "
    "ONLY if burden is near-zero and the tool lives inside the channel they already use (the chat). They will "
    "not run a separate structured system under load. Still to retrieve: a grassroots SUCCESS case (White "
    "Helmets local recording rescues; a camp committee with a real register) + a hard-collapse case + a "
    "coordinator interview answering the 6 questions.",
    "UPDATED: A6-EVIDENCE-CAMPAIGN.md (added 'Grassroots correction' section; split status into formal=Early "
    "Support vs grassroots=Unproven-Leaning-Weak) and EVIDENCE-CONVERGENCE-SCOREBOARD.md (A6 row now shows "
    "both). A1 preserved at bounded conclusion. NO DESIGN WORK.",
])
add_para(
    "A6 verdict after grassroots correction: formal orgs = Early Support; Nidaa-target grassroots coordinator "
    "= Unproven, leaning Weak. The project's real next question is no longer 'can structured info be "
    "maintained?' but 'can the kind of coordinator Nidaa depends on maintain it?'"
)

# ===== 14.23 — 2026-07-15 (evening): A6 split A6a/A6b; pause public expansion; field validation =====
add_heading("14.23 — 2026-07-15: A6 split into A6a/A6b; public expansion paused; A6b field validation", 2)
add_para(
    "Milestone finding recorded. The most important result of the project so far is NOT A1/A3 — it is that "
    "grassroots coordinator structured-maintenance is unproven/weak. Public evidence now has diminishing "
    "returns for A6b; the remaining gap is a real coordinator. No design/solution work."
)
add_bullets([
    "ASSUMPTION SPLIT (user direction): A6 -> A6a (formal orgs maintain structured info = Early Support) and "
    "A6b (grassroots coordinators maintain structured info = Weak/Unproven, the Nidaa user). A6b is the "
    "existential one. Scoreboard + CURRENT-EVIDENCE-STATE.md updated to show both.",
    "DEFINING FINDING locked in: 'They do coordinate, but via chats and committees, and structured/verified "
    "documentation is a gap, not a habit.' To be tested against real coordinators — if it survives, it changes "
    "what problem Nidaa must solve.",
    "PUBLIC-SOURCE EXPANSION FOR A6 PAUSED. No more reports/contradictions until a field response arrives.",
    "NEW FILE pilot/A6-FIELD-LOG.md — the home for the next milestone: ONE real coordinator workflow. Contains "
    "the Six A6 Questions as a standing processing rule: every incoming coordinator conversation (Omar, Shaima, "
    "Abdelrahman, White Helmets, ERRs, Sameer, iMMAP/referrals) is routed through the 6 questions and logged. "
    "Watch-list pre-loaded with pending conversations; success criterion = one six-answer coordinator story.",
    "NO DESIGN WORK. A1 preserved at bounded conclusion. A6b is now the highest-risk row and the active gate; "
    "its resolution depends on field evidence, not documents.",
])
add_para(
    "Next milestone is not another report. Next milestone is one real coordinator workflow, processed through "
    "the Six A6 Questions in A6-FIELD-LOG.md."
)

# ===== footer note =====
doc.add_paragraph()
fn = doc.add_paragraph()
fr = fn.add_run(
    "Build verification: npm run build passes; runtime API verified (GET/POST/verify, Arabic UTF-8 "
    "round-trip confirmed). Repo: github.com/theabdlrah/nidaa."
)
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = MUTED
frPr = fr._element.get_or_add_rPr()
frPr.append(OxmlElement("w:rtl"))

out = r"C:\Users\theab\Desktop\Nidaa-Journal.docx"
doc.save(out)
print("saved:", out)
